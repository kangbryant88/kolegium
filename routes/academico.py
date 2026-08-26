from flask import Blueprint, render_template, request, redirect, url_for, session, flash, send_file, make_response, jsonify, current_app
from datetime import datetime, date, timedelta
import io
import os
import re
import unicodedata
import pandas as pd
import openpyxl
from docx import Document
from fpdf import FPDF
import urllib.parse
import uuid
from models import db, Bitacora, Grado, Tema, AsistenciaDiaria, AsistenciaPersonal, Estudiante, Representante, Incidencia, AsistenciaEstudiante, AlertaDefensoria, EnlaceTemporal, SolicitudEnlace, SolicitudActualizacion, ProyectoAula, BancoIndicador, EvaluacionEstudiante, ProyectoAprendizaje, ProyectoArea, ProyectoEvaluacion

academico_bp = Blueprint('academico', __name__, url_prefix='/academico')

# ==========================================
# --- 4. PLANIFICADOR DOCENTE ---
# ==========================================

@academico_bp.route('/planificador')
def planificador():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    uid = session['usuario_id']
    clases = Bitacora.query.filter_by(usuario_id=uid).order_by(Bitacora.fecha.desc()).all()
    grados = Grado.query.all(); temas = Tema.query.all()
    return render_template('planificador.html', registros=clases, grados=grados, temas=temas,
                           total=len(clases), completadas=len([c for c in clases if c.estado == 'Completado']),
                           pendientes=len([c for c in clases if c.estado == 'Pendiente']))

@academico_bp.route('/agregar', methods=['POST'])
def agregar():
    hora = request.form.get('hora', '07:00')
    f = datetime.strptime(f"{datetime.now().strftime('%Y-%m-%d')} {hora}", "%Y-%m-%d %H:%M")
    db.session.add(Bitacora(fecha=f, grado=request.form['grado'], actividad=request.form['actividad'], 
                            estado=request.form['estado'], usuario_id=session['usuario_id']))
    db.session.commit(); return redirect(url_for('academico.planificador'))

@academico_bp.route('/editar/<int:id>', methods=['GET', 'POST'])
def editar(id):
    reg = Bitacora.query.get_or_404(id)
    if request.method == 'POST':
        reg.grado, reg.actividad, reg.estado = request.form['grado'], request.form['actividad'], request.form['estado']
        reg.fecha = datetime.strptime(f"{reg.fecha.strftime('%Y-%m-%d')} {request.form['hora']}", "%Y-%m-%d %H:%M")
        db.session.commit(); return redirect(url_for('academico.planificador'))
    return render_template('editar.html', registro=reg, grados=Grado.query.all())

@academico_bp.route('/eliminar/<int:id>')
def eliminar(id):
    db.session.delete(Bitacora.query.get_or_404(id)); db.session.commit()
    return redirect(url_for('academico.planificador'))

@academico_bp.route('/reporte_diario')
def reporte_diario():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    hoy = datetime.now().date()
    regs = [r for r in Bitacora.query.filter_by(usuario_id=session['usuario_id']).all() if r.fecha.date() == hoy]
    return render_template('reporte.html', registros=regs)

@academico_bp.route('/reporte_general')
def reporte_general():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    regs = Bitacora.query.filter_by(usuario_id=session['usuario_id']).order_by(Bitacora.fecha.desc()).all()
    return render_template('reporte.html', registros=regs)

# ==========================================
# --- 5. EXPORTACIONES (EXCEL / WORD) ---
# ==========================================

@academico_bp.route('/exportar_excel')
def exportar_excel():
    regs = Bitacora.query.filter_by(usuario_id=session['usuario_id']).all()
    datos = [{'Fecha': r.fecha.strftime('%d/%m/%Y'), 'Hora': r.fecha.strftime('%I:%M %p'), 
              'Grado': r.grado, 'Actividad': r.actividad, 'Estado': r.estado} for r in regs]
    df = pd.DataFrame(datos); output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Reporte', startrow=4, index=False)
        ws = writer.sheets['Reporte']; fmt = writer.book.add_format({'bold': True, 'font_size': 14})
        ws.write('A1', 'REPORTE DE ACTIVIDADES - EDUPALNNER OS', fmt)
        ws.write('A2', f"Docente: {session.get('nombre_completo')}")
        ws.write('A3', f"Área: {session.get('area_trabajo')}")
    output.seek(0)
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 
                     download_name="Reporte_EduPlanner.xlsx", as_attachment=True)

@academico_bp.route('/exportar_word')
def exportar_word():
    doc = Document(); doc.add_heading('REPORTE OFICIAL DE ACTIVIDADES', 0)
    p = doc.add_paragraph(); p.add_run('Docente: ').bold = True; p.add_run(f"{session.get('nombre_completo')}\n")
    p.add_run('Área: ').bold = True; p.add_run(f"{session.get('area_trabajo')}")
    table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
    for i, t in enumerate(['Fecha', 'Hora', 'Grado', 'Actividad', 'Estado']): table.rows[0].cells[i].text = t
    for r in Bitacora.query.filter_by(usuario_id=session['usuario_id']).all():
        row = table.add_row().cells
        row[0].text, row[1].text = r.fecha.strftime('%d/%m/%Y'), r.fecha.strftime('%I:%M %p')
        row[2].text, row[3].text, row[4].text = r.grado, r.actividad, r.estado
    f = io.BytesIO(); doc.save(f); f.seek(0)
    return send_file(f, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                     download_name="Reporte_EduPlanner.docx", as_attachment=True)

# ==========================================
# ==========================================
# --- 7. CONTROL DE ASISTENCIA ---
# ==========================================

@academico_bp.route('/asistencia', methods=['GET', 'POST'])
def asistencia():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    if request.method == 'POST':
        grado = Grado.query.get(request.form['grado_id'])
        v, h = int(request.form['varones'] or 0), int(request.form['hembras'] or 0)
        t = grado.total_varones + grado.total_hembras
        p = round(((v + h) / t) * 100, 1) if t > 0 else 0
        db.session.add(AsistenciaDiaria(fecha=datetime.strptime(request.form['fecha'], '%Y-%m-%d').date(), 
                       grado_seccion=grado.nombre, matricula_total=t, varones=v, hembras=h, asistentes=v+h, 
                       porcentaje=p, usuario_id=session['usuario_id']))
        db.session.commit(); return redirect(url_for('academico.asistencia'))
    if session.get('rol_id') in [1, 2]:
        regs = AsistenciaDiaria.query.order_by(AsistenciaDiaria.fecha.desc()).all()
    else:
        regs = AsistenciaDiaria.query.filter_by(usuario_id=session['usuario_id']).order_by(AsistenciaDiaria.fecha.desc()).all()
        
    from collections import OrderedDict
    
    meses_es = {1:'Enero', 2:'Febrero', 3:'Marzo', 4:'Abril', 5:'Mayo', 6:'Junio', 
                7:'Julio', 8:'Agosto', 9:'Septiembre', 10:'Octubre', 11:'Noviembre', 12:'Diciembre'}
    dias_es = {0:'Lunes', 1:'Martes', 2:'Miércoles', 3:'Jueves', 4:'Viernes', 5:'Sábado', 6:'Domingo'}

    def formato_fecha_es(fecha_obj):
        return f"{dias_es[fecha_obj.weekday()]}, {fecha_obj.day} de {meses_es[fecha_obj.month]} de {fecha_obj.year}"

    semanas = OrderedDict()
    
    for r in regs:
        dt = r.fecha
        lunes = dt - timedelta(days=dt.weekday())
        viernes = lunes + timedelta(days=4)
        semana_key = f"Semana del {lunes.strftime('%d/%m/%Y')} al {viernes.strftime('%d/%m/%Y')}"
        
        if semana_key not in semanas:
            semanas[semana_key] = OrderedDict()
            
        fecha_str = formato_fecha_es(dt)
        if fecha_str not in semanas[semana_key]:
            semanas[semana_key][fecha_str] = []
            
        semanas[semana_key][fecha_str].append(r)

    return render_template('asistencia.html', registros=regs, semanas=semanas, grados=Grado.query.all(), hoy=datetime.now().strftime('%Y-%m-%d'))

@academico_bp.route('/eliminar_asistencia/<int:id>', methods=['POST'])
def eliminar_asistencia(id):
    db.session.delete(AsistenciaDiaria.query.get_or_404(id)); db.session.commit()
    return redirect(url_for('academico.asistencia'))

@academico_bp.route('/asistencia_personal', methods=['GET', 'POST'])
def asistencia_personal():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    if request.method == 'POST':
        mat_base = int(request.form.get('matricula_base') or 0)
        asistentes = int(request.form.get('asistentes') or 0)
        pct = round((asistentes / mat_base) * 100, 1) if mat_base > 0 else 0.0
        db.session.add(AsistenciaPersonal(
            fecha=datetime.strptime(request.form['fecha'], '%Y-%m-%d').date(), 
            categoria=request.form['categoria'], 
            matricula_base=mat_base, 
            asistentes=asistentes, 
            porcentaje=pct, 
            usuario_id=session['usuario_id']
        ))
        db.session.commit()
        return redirect(url_for('academico.asistencia_personal'))
    if session.get('rol_id') in [1, 2]:
        regs = AsistenciaPersonal.query.order_by(AsistenciaPersonal.fecha.desc()).all()
    else:
        regs = AsistenciaPersonal.query.filter_by(usuario_id=session['usuario_id']).order_by(AsistenciaPersonal.fecha.desc()).all()
    
    # Memoria: Obtener el último valor registrado por categoría
    bases_conocidas = {}
    categorias = ['Docentes', 'Administrativos', 'Obreros', 'Vigilantes', 'Cocina']
    for cat in categorias:
        ultimo = AsistenciaPersonal.query.filter_by(usuario_id=session['usuario_id'], categoria=cat)\
            .order_by(AsistenciaPersonal.fecha.desc(), AsistenciaPersonal.id.desc()).first()
        if ultimo:
            bases_conocidas[cat] = ultimo.matricula_base

    return render_template('asistencia_personal.html', registros=regs, hoy=datetime.now().strftime('%Y-%m-%d'), bases=bases_conocidas)

@academico_bp.route('/eliminar_asistencia_personal/<int:id>')
def eliminar_asistencia_personal(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    db.session.delete(AsistenciaPersonal.query.get_or_404(id))
    db.session.commit()
    return redirect(url_for('academico.asistencia_personal'))

@academico_bp.route('/editar_asistencia_personal/<int:id>', methods=['GET', 'POST'])
def editar_asistencia_personal(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    reg = AsistenciaPersonal.query.get_or_404(id)
    if request.method == 'POST':
        mat_base = int(request.form.get('matricula_base') or 0)
        asistentes = int(request.form.get('asistentes') or 0)
        
        reg.fecha = datetime.strptime(request.form['fecha'], '%Y-%m-%d').date()
        reg.categoria = request.form['categoria']
        reg.matricula_base = mat_base
        reg.asistentes = asistentes
        reg.porcentaje = round((asistentes / mat_base) * 100, 1) if mat_base > 0 else 0.0
        
        db.session.commit()
        return redirect(url_for('academico.asistencia_personal'))
    
    return render_template('editar_asistencia_personal.html', registro=reg)

# ==========================================
# --- 8. CONFIGURACIONES GLOBALES ---
# ==========================================

@academico_bp.route('/historial_global')
def historial_global():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    tiene_acceso = (
        session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)'] or 
        (session.get('nombre_rol') == 'Administrativo' and session.get('departamento_asignado') == 'Dirección')
    )
    if not tiene_acceso:
        return redirect(url_for('index'))
    
    fecha_str = request.args.get('fecha')
    if fecha_str:
        try:
            fecha_busqueda = datetime.strptime(fecha_str, '%Y-%m-%d').date()
        except ValueError:
            fecha_busqueda = date.today() - timedelta(days=1)
    else:
        fecha_busqueda = date.today() - timedelta(days=1)
        
    regs_estudiantes = AsistenciaDiaria.query.filter_by(fecha=fecha_busqueda).all()
    regs_personal = AsistenciaPersonal.query.filter_by(fecha=fecha_busqueda).all()
    
    t_estudiantes = sum(r.asistentes for r in regs_estudiantes)
    t_personal = sum(r.asistentes for r in regs_personal)
    
    return render_template('historial.html', 
                           fecha_busqueda=fecha_busqueda.strftime('%Y-%m-%d'),
                           regs_estudiantes=regs_estudiantes,
                           regs_personal=regs_personal,
                           t_estudiantes=t_estudiantes,
                           t_personal=t_personal)

# ==========================================
# --- 10. MÓDULO DE ESTADÍSTICA ---
# ==========================================

def generar_cedula_escolar(nro_parto, anio_nino, cedula_rep):
    # [Nro Parto] + [Últimos 2 dígitos Año del Niño] + [Cédula Madre/Representante]
    anio_str = str(anio_nino)[-2:]
    return f"{str(nro_parto)}{anio_str}{str(cedula_rep)}"

@academico_bp.route('/estadistica_global')
def estadistica_global():
    if not session.get('logeado'):
        return redirect(url_for('auth.login'))
        
    estudiantes = Estudiante.query.all() or []
    total_estudiantes = len(estudiantes)
    ultimos_ingresos = Estudiante.query.filter_by(estatus='Activo').order_by(Estudiante.fecha_registro.desc()).limit(3).all() or []
    ultimos_egresos = Estudiante.query.filter_by(estatus='Egreso').order_by(Estudiante.fecha_registro.desc()).limit(3).all() or []
    
    return render_template('estadistica.html', 
                           estudiantes=estudiantes, 
                           total_estudiantes=total_estudiantes,
                           ultimos_ingresos=ultimos_ingresos,
                           ultimos_egresos=ultimos_egresos,
                           grados=Grado.query.all() or [])

@academico_bp.route('/api/estudiantes/<int:grado_id>')
def api_estudiantes_por_grado(grado_id):
    """API ligera: devuelve estudiantes de un grado en JSON para el modal dinámico."""
    if not session.get('logeado'):
        return jsonify({'error': 'No autorizado'}), 401

    grado = Grado.query.get_or_404(grado_id)
    estudiantes = Estudiante.query.filter_by(grado_id=grado_id).order_by(Estudiante.nombre_completo.asc()).all()
    
    resultado = {
        'grado_nombre': grado.nombre,
        'total': len(estudiantes),
        'estudiantes': [{
            'id': est.id,
            'cedula_escolar': est.cedula_escolar,
            'nombre_completo': est.nombre_completo,
            'inicial': est.nombre_completo[0].upper() if est.nombre_completo else '?',
            'genero': est.genero or 'No especificado',
            'estatus': est.estatus or 'Activo',
            'cedula_rep': est.representante_info.cedula if est.representante_info else 'N/A'
        } for est in estudiantes]
    }
    return jsonify(resultado)

@academico_bp.route('/registrar_estudiante', methods=['POST'])
def registrar_estudiante():
    if not session.get('logeado'):
        return redirect(url_for('auth.login'))

    # Obtener datos del representante
    cedula_rep = request.form.get('cedula_representante')
    
    # Verificar si el representante ya existe para no duplicarlo
    representante = Representante.query.filter_by(cedula=cedula_rep).first()
    if not representante:
        representante = Representante(
            cedula=cedula_rep,
            nombre_completo=request.form.get('nombre_representante'),
            telefono=request.form.get('telefono_representante'),
            email=request.form.get('email_representante'),
            email_representante=request.form.get('email_representante'),
            parentesco=request.form.get('parentesco'),
            direccion_habitacion=request.form.get('direccion_habitacion'),
            direccion_completa=request.form.get('direccion_completa'),
            ocupacion=request.form.get('ocupacion_representante'),
            lugar_direccion_trabajo=request.form.get('lugar_direccion_trabajo'),
            banco_nombre=request.form.get('banco_nombre'),
            banco_cuenta_numero=request.form.get('banco_cuenta_numero'),
            banco_cuenta_tipo=request.form.get('banco_cuenta_tipo'),
            banco_titular_nombre=request.form.get('banco_titular_nombre'),
            banco_titular_ci=request.form.get('banco_titular_ci')
        )
        db.session.add(representante)
        db.session.flush() # Para poder usar el id del representante
    else:
        # Update if not set or just leave it
        if request.form.get('direccion_habitacion'):
            representante.direccion_habitacion = request.form.get('direccion_habitacion')
        if request.form.get('direccion_completa'):
            representante.direccion_completa = request.form.get('direccion_completa')
        if request.form.get('email_representante'):
            representante.email = request.form.get('email_representante')
            representante.email_representante = request.form.get('email_representante')
        if request.form.get('ocupacion_representante'):
            representante.ocupacion = request.form.get('ocupacion_representante')
        if request.form.get('lugar_direccion_trabajo'):
            representante.lugar_direccion_trabajo = request.form.get('lugar_direccion_trabajo')
        if request.form.get('banco_nombre'):
            representante.banco_nombre = request.form.get('banco_nombre')
        if request.form.get('banco_cuenta_numero'):
            representante.banco_cuenta_numero = request.form.get('banco_cuenta_numero')
        if request.form.get('banco_cuenta_tipo'):
            representante.banco_cuenta_tipo = request.form.get('banco_cuenta_tipo')
        if request.form.get('banco_titular_nombre'):
            representante.banco_titular_nombre = request.form.get('banco_titular_nombre')
        if request.form.get('banco_titular_ci'):
            representante.banco_titular_ci = request.form.get('banco_titular_ci')
        db.session.flush()

    # Calcular datos para generar cédula escolar
    nro_parto = request.form.get('nro_parto', '1')
    fecha_nac_str = request.form.get('fecha_nacimiento')
    fecha_nac = datetime.strptime(fecha_nac_str, '%Y-%m-%d').date()
    anio_nino = fecha_nac.year

    cedula_escolar_gen = generar_cedula_escolar(nro_parto, anio_nino, cedula_rep)

    estudiante_existente = Estudiante.query.filter_by(cedula_escolar=cedula_escolar_gen).first()
    if estudiante_existente:
        grado_nombre = estudiante_existente.grado.nombre if estudiante_existente.grado else "el sistema"
        flash(f'Error: Este niño ya ha sido registrado anteriormente y se encuentra en {grado_nombre}.', 'error')
        return redirect(url_for('academico.estadistica_global'))

    # Valores numéricos o nulos
    talla_val = request.form.get('talla')
    peso_val = request.form.get('peso')
    edad_val = request.form.get('edad')
    nombres_val = request.form.get('nombres', '')
    apellidos_val = request.form.get('apellidos', '')

    # Guardar Estudiante
    estudiante = Estudiante(
        cedula_escolar=cedula_escolar_gen,
        nombre_completo=f"{nombres_val} {apellidos_val}".strip(),
        nombres=nombres_val,
        apellidos=apellidos_val,
        fecha_nacimiento=fecha_nac,
        lugar_nacimiento=request.form.get('lugar_nacimiento'),
        municipio=request.form.get('municipio'),
        parroquia=request.form.get('parroquia'),
        edad=int(edad_val) if edad_val else None,
        genero=request.form.get('genero'),
        nacionalidad=request.form.get('nacionalidad', 'Venezolana'),
        num_acta=request.form.get('num_acta'),
        num_oficio=request.form.get('num_oficio'),
        direccion_alumno=request.form.get('direccion_alumno'),
        telefono_habitacion=request.form.get('telefono_habitacion'),
        posee_canaima=True if request.form.get('posee_canaima') == 'on' else False,
        talla=float(talla_val) if talla_val else None,
        peso=float(peso_val) if peso_val else None,
        calzado=request.form.get('calzado'),
        talla_camisa=request.form.get('talla_camisa'),
        talla_pantalon=request.form.get('talla_pantalon'),
        tipaje=request.form.get('tipaje'),
        vacunacion_completa=request.form.get('vacunacion_completa'),
        alergias=request.form.get('alergias'),
        neurodivergencia=True if request.form.get('neurodivergencia') == 'on' else False,
        neuro_detalle=request.form.get('neuro_detalle'),
        posee_enfermedad=True if request.form.get('posee_enfermedad') == 'on' else False,
        enfermedad_detalle=request.form.get('enfermedad_detalle'),
        toma_medicamento=True if request.form.get('toma_medicamento') == 'on' else False,
        medicamento_detalle=request.form.get('medicamento_detalle'),
        alergico_medicamento=True if request.form.get('alergico_medicamento') == 'on' else False,
        alergia_medicamento_detalle=request.form.get('alergia_medicamento_detalle'),
        literal=request.form.get('literal_escolar'),
        literal_escolar=request.form.get('literal_escolar'),
        procedencia=request.form.get('plantel_procedencia'),
        plantel_procedencia=request.form.get('plantel_procedencia'),
        email_estudiante=request.form.get('email_estudiante'),
        es_repetidor=True if request.form.get('es_repetidor') == 'on' else False,
        doc_partida=True if request.form.get('doc_partida') == 'on' else False,
        doc_sano=True if request.form.get('doc_sano') == 'on' else False,
        doc_vacuna=True if request.form.get('doc_vacuna') == 'on' else False,
        lateralidad=request.form.get('lateralidad'),
        nuevo_ingreso=True if request.form.get('nuevo_ingreso') == 'on' else False,
        institucion_procedencia=request.form.get('institucion_procedencia'),
        estatus=request.form.get('estatus', 'Activo'),
        madre_nombre=request.form.get('madre_nombre'),
        madre_ci=request.form.get('madre_ci'),
        madre_ocupacion=request.form.get('madre_ocupacion'),
        madre_telefono=request.form.get('madre_telefono'),
        madre_direccion=request.form.get('madre_direccion'),
        padre_nombre=request.form.get('padre_nombre'),
        padre_ci=request.form.get('padre_ci'),
        padre_ocupacion=request.form.get('padre_ocupacion'),
        padre_telefono=request.form.get('padre_telefono'),
        padre_direccion=request.form.get('padre_direccion'),
        telefono_familiar_extra=request.form.get('telefono_familiar_extra'),
        autorizacion_odontologica=True if request.form.get('autorizacion_odontologica') == 'on' else False,
        grado_id=request.form.get('grado_id'),
        representante_id=representante.id
    )

    try:
        db.session.add(estudiante)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        flash('Ocurrió un error inesperado al guardar el estudiante. Inténtalo de nuevo.', 'error')
    
    return redirect(url_for('academico.estadistica_global'))

@academico_bp.route('/perfil_estudiante/<int:id>')
def perfil_estudiante(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    estudiante = Estudiante.query.get_or_404(id)
    return render_template('perfil_estudiante.html', estudiante=estudiante, grados=Grado.query.all())

@academico_bp.route('/editar_estudiante/<int:id>', methods=['POST'])
def editar_estudiante(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    est = Estudiante.query.get_or_404(id)

    nombres_val = request.form.get('nombres', '')
    apellidos_val = request.form.get('apellidos', '')
    edad_val = request.form.get('edad')

    est.nombres = nombres_val
    est.apellidos = apellidos_val
    est.nombre_completo = f"{nombres_val} {apellidos_val}".strip()
    est.fecha_nacimiento = datetime.strptime(request.form.get('fecha_nacimiento'), '%Y-%m-%d').date()
    est.lugar_nacimiento = request.form.get('lugar_nacimiento')
    est.municipio = request.form.get('municipio')
    est.parroquia = request.form.get('parroquia')
    est.edad = int(edad_val) if edad_val else None
    est.genero = request.form.get('genero')
    est.nacionalidad = request.form.get('nacionalidad', 'Venezolana')
    est.num_acta = request.form.get('num_acta')
    est.num_oficio = request.form.get('num_oficio')
    est.direccion_alumno = request.form.get('direccion_alumno')
    est.telefono_habitacion = request.form.get('telefono_habitacion')
    est.posee_canaima = True if request.form.get('posee_canaima') == 'on' else False

    talla_val = request.form.get('talla')
    peso_val = request.form.get('peso')
    est.talla = float(talla_val) if talla_val else None
    est.peso = float(peso_val) if peso_val else None

    est.calzado = request.form.get('calzado')
    est.talla_camisa = request.form.get('talla_camisa')
    est.talla_pantalon = request.form.get('talla_pantalon')
    est.tipaje = request.form.get('tipaje')
    est.vacunacion_completa = request.form.get('vacunacion_completa')
    est.alergias = request.form.get('alergias')
    est.neurodivergencia = True if request.form.get('neurodivergencia') == 'on' else False
    est.neuro_detalle = request.form.get('neuro_detalle')
    est.posee_enfermedad = True if request.form.get('posee_enfermedad') == 'on' else False
    est.enfermedad_detalle = request.form.get('enfermedad_detalle')
    est.toma_medicamento = True if request.form.get('toma_medicamento') == 'on' else False
    est.medicamento_detalle = request.form.get('medicamento_detalle')
    est.alergico_medicamento = True if request.form.get('alergico_medicamento') == 'on' else False
    est.alergia_medicamento_detalle = request.form.get('alergia_medicamento_detalle')
    est.literal = request.form.get('literal_escolar')
    est.literal_escolar = request.form.get('literal_escolar')
    est.procedencia = request.form.get('plantel_procedencia')
    est.plantel_procedencia = request.form.get('plantel_procedencia')
    est.email_estudiante = request.form.get('email_estudiante')
    est.es_repetidor = True if request.form.get('es_repetidor') == 'on' else False
    est.doc_partida = True if request.form.get('doc_partida') == 'on' else False
    est.doc_sano = True if request.form.get('doc_sano') == 'on' else False
    est.doc_vacuna = True if request.form.get('doc_vacuna') == 'on' else False
    est.lateralidad = request.form.get('lateralidad')
    est.nuevo_ingreso = True if request.form.get('nuevo_ingreso') == 'on' else False
    est.institucion_procedencia = request.form.get('institucion_procedencia')
    est.grado_id = request.form.get('grado_id')

    # Datos de la Madre
    est.madre_nombre = request.form.get('madre_nombre')
    est.madre_ci = request.form.get('madre_ci')
    est.madre_ocupacion = request.form.get('madre_ocupacion')
    est.madre_telefono = request.form.get('madre_telefono')
    est.madre_direccion = request.form.get('madre_direccion')
    # Datos del Padre
    est.padre_nombre = request.form.get('padre_nombre')
    est.padre_ci = request.form.get('padre_ci')
    est.padre_ocupacion = request.form.get('padre_ocupacion')
    est.padre_telefono = request.form.get('padre_telefono')
    est.padre_direccion = request.form.get('padre_direccion')
    # Otros y Autorizaciones
    est.telefono_familiar_extra = request.form.get('telefono_familiar_extra')
    est.autorizacion_odontologica = True if request.form.get('autorizacion_odontologica') == 'on' else False

    # Actualizar (en el mismo registro, sin reasignar) al Representante vinculado
    if est.representante_info:
        rep = est.representante_info
        rep.cedula = request.form.get('cedula_representante', rep.cedula)
        rep.nombre_completo = request.form.get('nombre_representante', rep.nombre_completo)
        rep.telefono = request.form.get('telefono_representante')
        rep.parentesco = request.form.get('parentesco')
        rep.email = request.form.get('email_representante')
        rep.email_representante = request.form.get('email_representante')
        rep.direccion_habitacion = request.form.get('direccion_habitacion')
        rep.direccion_completa = request.form.get('direccion_completa')
        rep.ocupacion = request.form.get('ocupacion_representante')
        rep.lugar_direccion_trabajo = request.form.get('lugar_direccion_trabajo')
        rep.banco_nombre = request.form.get('banco_nombre')
        rep.banco_cuenta_numero = request.form.get('banco_cuenta_numero')
        rep.banco_cuenta_tipo = request.form.get('banco_cuenta_tipo')
        rep.banco_titular_nombre = request.form.get('banco_titular_nombre')
        rep.banco_titular_ci = request.form.get('banco_titular_ci')

    db.session.commit()
    return redirect(url_for('academico.perfil_estudiante', id=est.id))

@academico_bp.route('/egresar_estudiante/<int:id>', methods=['POST'])
def egresar_estudiante(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    est = Estudiante.query.get_or_404(id)
    est.estatus = 'Activo' if est.estatus == 'Egreso' else 'Egreso'
    db.session.commit()
    return redirect(url_for('academico.perfil_estudiante', id=est.id))

@academico_bp.route('/eliminar_estudiante/<int:id>', methods=['POST'])
def eliminar_estudiante(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    est = Estudiante.query.get_or_404(id)
    db.session.delete(est)
    db.session.commit()
    return redirect(url_for('academico.estadistica_global'))

# ==========================================
# --- 11. MÓDULO MI AULA ---
# ==========================================

@academico_bp.route('/mi_aula', methods=['GET', 'POST'])
def mi_aula():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    
    rol = session.get('nombre_rol')
    if rol not in ['Administrador Supremo', 'Equipo Directivo (Dirección)', 'Docente de Aula']:
        return redirect(url_for('index'))
        
    grados = []
    grado_seleccionado = None
    estudiantes = []
    asistencia_porcentaje = {}
    grado_id = request.args.get('grado_id')

    if rol == 'Docente de Aula':
        if not grado_id:
            usuario_id = session.get('usuario_id')
            salon = Grado.query.filter(Grado.docentes.any(id=usuario_id)).first()
            if salon:
                return redirect(url_for('academico.mi_aula', grado_id=salon.id))
            else:
                flash("Aún no tienes un grado asignado.", "warning")
                return redirect(url_for('index'))
        else:
            grado_seleccionado = Grado.query.get(grado_id)
            if grado_seleccionado and session.get('usuario_id') in [d.id for d in grado_seleccionado.docentes]:
                estudiantes = Estudiante.query.filter_by(grado_id=grado_seleccionado.id, estatus='Activo').order_by(Estudiante.nombre_completo.asc()).all()
            else:
                flash("No tienes acceso a este salón.", "danger")
                return redirect(url_for('index'))
    else:
        grados = Grado.query.all()
        if grado_id:
            grado_seleccionado = Grado.query.get(grado_id)
            if grado_seleccionado:
                estudiantes = Estudiante.query.filter_by(grado_id=grado_seleccionado.id, estatus='Activo').order_by(Estudiante.nombre_completo.asc()).all()

    for est in estudiantes:
        total_dias = AsistenciaEstudiante.query.filter_by(estudiante_id=est.id).count() or 0
        asistencias = AsistenciaEstudiante.query.filter_by(estudiante_id=est.id, estatus='Presente').count() or 0
        porcentaje = (asistencias / total_dias * 100) if total_dias > 0 else 0.0
        asistencia_porcentaje[est.id] = round(porcentaje, 1)
        
    datos_incidencias = {'Conducta': 0, 'Académico': 0, 'Salud': 0, 'Familiar': 0}
    asistencia_promedio_salon = 100.0
    total_matricula = len(estudiantes)

    if estudiantes:
        estudiante_ids = [e.id for e in estudiantes]
        mes_actual = date.today().month
        
        incidencias_salon = Incidencia.query.filter(Incidencia.estudiante_id.in_(estudiante_ids)).all()
        for inc in incidencias_salon:
            if inc.fecha and inc.fecha.month == mes_actual and inc.categoria in datos_incidencias:
                datos_incidencias[inc.categoria] += 1
                
        treinta_dias_atras = date.today() - timedelta(days=30)
        asist_records = AsistenciaEstudiante.query.filter(
            AsistenciaEstudiante.estudiante_id.in_(estudiante_ids),
            AsistenciaEstudiante.fecha >= treinta_dias_atras
        ).all()
        
        if asist_records:
            total_dias_salon = len(asist_records)
            asistencias_positivas = sum(1 for a in asist_records if a.estatus == 'Presente')
            asistencia_promedio_salon = round((asistencias_positivas / total_dias_salon) * 100, 1) if total_dias_salon > 0 else 0.0
            
    total_varones = sum(1 for e in estudiantes if e.genero == 'Masculino')
    total_hembras = sum(1 for e in estudiantes if e.genero == 'Femenino')
    docente_titular = ", ".join([d.nombre_completo for d in grado_seleccionado.docentes]) if (grado_seleccionado and grado_seleccionado.docentes) else "Docente no asignado"

    # Evaluación Descriptiva: Proyecto de Aula por momento (del salón) y Banco de Indicadores (del docente logeado)
    proyectos_aula = {}
    banco_indicadores = []
    if grado_seleccionado:
        for p in ProyectoAula.query.filter_by(grado_id=grado_seleccionado.id).all():
            proyectos_aula[p.momento] = p
    if session.get('usuario_id'):
        banco_indicadores = BancoIndicador.query.filter_by(docente_id=session['usuario_id']) \
            .order_by(BancoIndicador.momento.asc(), BancoIndicador.nivel.asc(), BancoIndicador.fecha_creacion.desc()).all()

    # Proyectos de Aprendizaje oficiales creados por el docente actual
    proyectos_oficiales = []
    if session.get('usuario_id'):
        proyectos_oficiales = ProyectoAprendizaje.query.filter_by(docente_id=session['usuario_id']) \
            .order_by(ProyectoAprendizaje.fecha_creacion.desc()).all()

    # Evaluaciones ya guardadas por estudiante, agrupadas por momento: {estudiante_id: {momento: EvaluacionEstudiante}}
    evaluaciones_por_estudiante = {}
    if estudiantes:
        evals = EvaluacionEstudiante.query.filter(
            EvaluacionEstudiante.estudiante_id.in_([e.id for e in estudiantes])
        ).all()
        for ev in evals:
            evaluaciones_por_estudiante.setdefault(ev.estudiante_id, {})[ev.momento] = ev

    estado_solicitudes = {}
    if estudiantes:
        mis_solicitudes = SolicitudEnlace.query.filter(
            SolicitudEnlace.estudiante_id.in_([e.id for e in estudiantes])
        ).order_by(SolicitudEnlace.fecha_solicitud.desc()).all()
        for sol in mis_solicitudes:
            if sol.estudiante_id not in estado_solicitudes:
                estado_solicitudes[sol.estudiante_id] = sol

    # Nuevas métricas: Gráfica mensual y Alertas de Defensoría
    datos_mensuales = {}
    alertas_activas = {}
    
    if estudiantes:
        hoy = date.today()
        # Obtener asistencias del mes actual para todos los estudiantes
        asistencias_mes = AsistenciaEstudiante.query.filter(
            AsistenciaEstudiante.estudiante_id.in_([e.id for e in estudiantes])
        ).all()
        
        for est in estudiantes:
            datos_mensuales[est.id] = {'Presente': 0, 'Ausente': 0, 'Justificado': 0}
            alertas_activas[est.id] = False
            
        for a in asistencias_mes:
            if a.fecha and a.fecha.month == hoy.month and a.fecha.year == hoy.year:
                if a.estatus in datos_mensuales[a.estudiante_id]:
                    datos_mensuales[a.estudiante_id][a.estatus] += 1
                    
        # Buscar Alertas de Defensoría activas (Pendiente) para estos estudiantes
        alertas = AlertaDefensoria.query.filter(
            AlertaDefensoria.estudiante_id.in_([e.id for e in estudiantes]),
            AlertaDefensoria.estatus_atencion == 'Pendiente'
        ).all()
        
        for alerta in alertas:
            # Simplificar el motivo para mostrarlo en el badge
            tipo_alerta = "Inasistencia" if alerta.motivo and "inasistencia" in alerta.motivo.lower() else "Incidencia"
            alertas_activas[alerta.estudiante_id] = {'tipo': tipo_alerta, 'id': alerta.id}

    return render_template('mi_aula.html', 
                           grado=grado_seleccionado, 
                           grados=grados, 
                           estudiantes=estudiantes, 
                           asistencia_porcentaje=asistencia_porcentaje,
                           datos_mensuales_estudiantes=datos_mensuales,
                           alertas_estudiantes=alertas_activas,
                           total_matricula=total_matricula,
                           total_varones=total_varones,
                           total_hembras=total_hembras,
                           docente_titular=docente_titular,
                           datos_incidencias=datos_incidencias,
                           asistencia_promedio_salon=asistencia_promedio_salon,
                           estado_solicitudes=estado_solicitudes,
                           proyectos_aula=proyectos_aula,
                           banco_indicadores=banco_indicadores,
                           proyectos_oficiales=proyectos_oficiales,
                           evaluaciones_por_estudiante=evaluaciones_por_estudiante)

@academico_bp.route('/guardar_asistencia_aula', methods=['POST'])
def guardar_asistencia_aula():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    grado_id = request.form.get('grado_id')
    estudiantes = Estudiante.query.filter_by(grado_id=grado_id).all()
    fecha_hoy = date.today()
    
    for est in estudiantes:
        vino = request.form.get(f'asistio_{est.id}') == 'on'
        registro = AsistenciaEstudiante.query.filter_by(estudiante_id=est.id, fecha=fecha_hoy).first()
        if registro:
            registro.estatus = 'Presente' if vino else 'Ausente'
            registro.grado_id = grado_id
        else:
            nuevo_registro = AsistenciaEstudiante(fecha=fecha_hoy, estatus='Presente' if vino else 'Ausente', estudiante_id=est.id, grado_id=grado_id)
            db.session.add(nuevo_registro)
            
    db.session.commit()
    url = url_for('academico.mi_aula')
    if session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)'] and grado_id:
        url = url_for('academico.mi_aula', grado_id=grado_id)
    return redirect(url)

@academico_bp.route('/agregar_incidencia', methods=['POST'])
def agregar_incidencia():
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    estudiante_id = request.form.get('estudiante_id')
    categoria = request.form.get('categoria')
    descripcion = request.form.get('descripcion')
    grado_id = request.form.get('grado_id')
    
    if estudiante_id and categoria and descripcion:
        incidencia = Incidencia(
            categoria=categoria,
            descripcion=descripcion,
            estudiante_id=estudiante_id,
            usuario_id=session.get('usuario_id')
        )
        db.session.add(incidencia)
        db.session.commit()
        
    url = url_for('academico.mi_aula')
    if session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)'] and grado_id:
        url = url_for('academico.mi_aula', grado_id=grado_id)
    return redirect(url)

@academico_bp.route('/descargar_inscripcion_inicial/<int:grado_id>')
def descargar_inscripcion_inicial(grado_id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    
    grado = Grado.query.get_or_404(grado_id)
    estudiantes = Estudiante.query.filter_by(grado_id=grado.id).order_by(Estudiante.nombre_completo.asc()).all()
    
    docente_nombre = ", ".join([d.nombre_completo for d in grado.docentes]) if grado.docentes else "No asignado"
    
    pdf = FPDF(orientation='L', unit='mm', format='Legal')
    pdf.add_page()
    
    # Membrete Oficial
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(330, 7, f"INSCRIPCIÓN INICIAL - {date.today().year}", ln=1, align='C')
    pdf.set_font('Arial', 'B', 9)
    pdf.cell(330, 5, f"Grado y Sección: {grado.nombre} | Docente: {docente_nombre}", ln=1, align='C')
    pdf.ln(5)
    
    pdf.set_font('Arial', 'B', 7)
    
    # Definición de anchos de las 19 columnas (~335mm disponibles, sumaremos 332)
    w = [6, 38, 16, 14, 8, 8, 24, 8, 38, 17, 18, 50, 45, 7, 7, 7, 7, 8, 8]
    
    # Calcular edad
    def calc_edad(fn):
        if not fn: return 'S/F'
        hoy = date.today()
        return str(hoy.year - fn.year - ((hoy.month, hoy.day) < (fn.month, fn.day)))
    
    headers = ["N°", "Apellidos y Nombres", "C.I. / C.E.", "F. Nac.", "Edad", "Sexo", "Procedencia", "Repet.", "Representante", "C.I. Rep.", "Teléfono", "Correo Representante", "Dirección Habitación", "Lit.", "Talla", "Peso", "Calz.", "T. Cam.", "T. Pan."]
    for i in range(len(headers)):
        pdf.cell(w[i], 6, headers[i], border=1, align='C')
    pdf.ln()
    
    def imprimir_celda_ajustada(ancho, texto, alineacion='L'):
        tamano_original = 7
        tamano_actual = tamano_original
        pdf.set_font_size(tamano_actual)
        while pdf.get_string_width(texto) > (ancho - 1) and tamano_actual > 3:
            tamano_actual -= 0.5
            pdf.set_font_size(tamano_actual)
        pdf.cell(ancho, 6, texto, border=1, align=alineacion)
        pdf.set_font_size(tamano_original)

    pdf.set_font('Arial', '', 7)
    for i, est in enumerate(estudiantes, 1):
        rep_nombre = est.representante_info.nombre_completo if est.representante_info else "Sin registro"
        rep_ci = est.representante_info.cedula if est.representante_info else "Sin registro"
        rep_tlf = est.representante_info.telefono if est.representante_info else "Sin registro"
        rep_email = est.representante_info.email_representante if est.representante_info and est.representante_info.email_representante else (est.representante_info.email if est.representante_info else "Sin registro")
        rep_dir = est.representante_info.direccion_completa if est.representante_info and est.representante_info.direccion_completa else (est.representante_info.direccion_habitacion if est.representante_info else "Sin registro")
        
        nombre_str = est.nombre_completo
        rep_str = rep_nombre
        email_str = rep_email
        dir_str = rep_dir
        proc = est.plantel_procedencia or est.procedencia or est.institucion_procedencia
        proc_str = proc if proc else "Ninguna"
        
        pdf.cell(w[0], 6, str(i), border=1, align='C')
        imprimir_celda_ajustada(w[1], nombre_str, 'L')
        pdf.cell(w[2], 6, str(est.cedula_escolar), border=1, align='C')
        pdf.cell(w[3], 6, est.fecha_nacimiento.strftime('%d/%m/%y') if est.fecha_nacimiento else 'S/F', border=1, align='C')
        pdf.cell(w[4], 6, calc_edad(est.fecha_nacimiento), border=1, align='C')
        pdf.cell(w[5], 6, "M" if est.genero == "Masculino" else ("F" if est.genero == "Femenino" else "-"), border=1, align='C')
        imprimir_celda_ajustada(w[6], proc_str, 'L')
        pdf.cell(w[7], 6, "Sí" if est.es_repetidor else "No", border=1, align='C')
        imprimir_celda_ajustada(w[8], rep_str, 'L')
        pdf.cell(w[9], 6, str(rep_ci), border=1, align='C')
        pdf.cell(w[10], 6, rep_tlf, border=1, align='C')
        imprimir_celda_ajustada(w[11], email_str, 'L')
        imprimir_celda_ajustada(w[12], dir_str, 'L')
        pdf.cell(w[13], 6, est.literal_escolar or est.literal or "-", border=1, align='C')
        pdf.cell(w[14], 6, f"{est.talla or '-'}", border=1, align='C')
        pdf.cell(w[15], 6, f"{est.peso or '-'}", border=1, align='C')
        pdf.cell(w[16], 6, str(est.calzado or '-'), border=1, align='C')
        pdf.cell(w[17], 6, str(est.talla_camisa or '-'), border=1, align='C')
        pdf.cell(w[18], 6, str(est.talla_pantalon or '-'), border=1, align='C')
        pdf.ln()

    from flask import make_response
    response = make_response(pdf.output(dest='S').encode('latin1'))
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=Inscripcion_Inicial_Grado_{grado.id}.pdf'
    return response

# ==========================================
# --- 11.1 EVALUACIÓN DESCRIPTIVA Y BOLETINES ---
# ==========================================

@academico_bp.route('/guardar_proyecto_aula', methods=['POST'])
def guardar_proyecto_aula():
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    grado_id = request.form.get('grado_id')
    momento_val = request.form.get('momento')
    titulo = request.form.get('titulo_proyecto', '').strip()

    if not grado_id or not momento_val or not titulo:
        flash('Debe indicar el salón, el momento y el título del Proyecto de Aula.', 'danger')
        return redirect(request.referrer or url_for('academico.mi_aula'))

    momento = int(momento_val)

    proyecto = ProyectoAula.query.filter_by(grado_id=grado_id, momento=momento).first()
    if not proyecto:
        proyecto = ProyectoAula(grado_id=grado_id, momento=momento)
        db.session.add(proyecto)

    proyecto.titulo_proyecto = titulo

    fecha_inicio = request.form.get('fecha_inicio')
    fecha_cierre = request.form.get('fecha_cierre')
    proyecto.fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date() if fecha_inicio else None
    proyecto.fecha_cierre = datetime.strptime(fecha_cierre, '%Y-%m-%d').date() if fecha_cierre else None

    db.session.commit()
    flash(f'Proyecto de Aula del Momento {momento} guardado correctamente.', 'success')
    return redirect(url_for('academico.mi_aula', grado_id=grado_id))


@academico_bp.route('/guardar_indicador', methods=['POST'])
def guardar_indicador():
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    texto = request.form.get('texto_indicador', '').strip()
    momento_val = request.form.get('momento')
    nivel = request.form.get('nivel')
    grado_id = request.form.get('grado_id')

    if not texto or not momento_val or not nivel:
        flash('Complete el momento, el nivel y el texto del indicador.', 'danger')
        return redirect(request.referrer or url_for('academico.mi_aula', grado_id=grado_id))

    indicador = BancoIndicador(
        docente_id=session['usuario_id'],
        momento=int(momento_val),
        nivel=nivel,
        texto_indicador=texto
    )
    db.session.add(indicador)
    db.session.commit()
    flash('Indicador agregado al banco.', 'success')
    return redirect(url_for('academico.mi_aula', grado_id=grado_id))


@academico_bp.route('/eliminar_indicador/<int:id>', methods=['POST'])
def eliminar_indicador(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    indicador = BancoIndicador.query.get_or_404(id)
    es_propietario = indicador.docente_id == session.get('usuario_id')
    es_admin = session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)']

    if not (es_propietario or es_admin):
        flash('No tiene permiso para eliminar este indicador.', 'danger')
        return redirect(request.referrer or url_for('academico.mi_aula'))

    grado_id = request.form.get('grado_id')
    db.session.delete(indicador)
    db.session.commit()
    flash('Indicador eliminado del banco.', 'success')
    return redirect(url_for('academico.mi_aula', grado_id=grado_id))


def _capitalizar_corto(texto):
    """
    Normaliza campos de texto corto (ej. Área de Formación, Componente) antes
    de guardarlos: quita espacios sobrantes y capitaliza cada palabra, para
    que "lenguaje Y comunicación" quede "Lenguaje Y Comunicación" -limpio y
    presentable en el PDF- sin importar cómo lo haya tipeado el docente.
    """
    return texto.strip().title() if texto else texto


def _volcar_datos_generales_proyecto(proyecto):
    """
    Copia los campos de las pestañas 1 y 3 del formulario (Datos Generales y
    Cierre) sobre `proyecto`. Se usa tanto al crear como al editar -en
    edición, `proyecto` ya existe y esto simplemente sobrescribe sus campos-.
    """
    fecha_inicio = request.form.get('fecha_inicio')
    fecha_culminacion = request.form.get('fecha_culminacion')
    fecha_entrega = request.form.get('fecha_entrega')

    proyecto.grado_id = request.form.get('grado_id')
    proyecto.tema = request.form.get('tema', '').strip()
    proyecto.momento_pedagogico = request.form.get('momento_pedagogico')
    proyecto.fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date() if fecha_inicio else None
    proyecto.fecha_culminacion = datetime.strptime(fecha_culminacion, '%Y-%m-%d').date() if fecha_culminacion else None
    proyecto.fecha_entrega = datetime.strptime(fecha_entrega, '%Y-%m-%d').date() if fecha_entrega else None
    proyecto.diagnostico_pedagogico = request.form.get('diagnostico_pedagogico')
    proyecto.proposito_integral = request.form.get('proposito_integral')
    proyecto.cierre_demostracion = request.form.get('cierre_demostracion')
    proyecto.cierre_dramatizacion = request.form.get('cierre_dramatizacion')
    proyecto.cierre_muestra = request.form.get('cierre_muestra')


def _reconstruir_areas_proyecto(proyecto):
    """
    Reemplaza por completo las Áreas de Formación (y sus Evaluaciones
    anidadas) de `proyecto` con lo enviado en la pestaña 2 del formulario.

    En edición, borrar y volver a crear es más simple y confiable que
    intentar emparejar cada fila existente con su índice del formulario
    -sobre todo porque el docente puede añadir/quitar bloques de área
    libremente en el cliente-. La cascada del modelo (`cascade='all,
    delete-orphan'`) se encarga de las evaluaciones asociadas.
    """
    for area_existente in list(proyecto.areas):
        db.session.delete(area_existente)
    db.session.flush()

    tema_indispensable_l = request.form.getlist('tema_indispensable[]')
    area_formacion_l = request.form.getlist('area_formacion[]')
    enfasis_l = request.form.getlist('enfasis[]')
    componente_l = request.form.getlist('componente[]')
    contenidos_l = request.form.getlist('contenidos[]')
    estrategia_pedagogica_l = request.form.getlist('estrategia_pedagogica[]')
    actividades_l = request.form.getlist('actividades[]')
    aprendizajes_esperados_l = request.form.getlist('aprendizajes_esperados[]')
    observaciones_l = request.form.getlist('observaciones[]')

    # --- Evaluación de cada área (misma posición/índice que la lista de áreas) ---
    estrategias_evaluacion_l = request.form.getlist('estrategias_evaluacion[]')
    indicadores_evaluar_l = request.form.getlist('indicadores_evaluar[]')
    tecnicas_evaluacion_l = request.form.getlist('tecnicas_evaluacion[]')
    instrumentos_evaluacion_l = request.form.getlist('instrumentos_evaluacion[]')
    tipos_evaluacion_l = request.form.getlist('tipos_evaluacion[]')
    formas_participacion_l = request.form.getlist('formas_participacion[]')
    recursos_materiales_l = request.form.getlist('recursos_materiales[]')

    def campo(lista, i):
        return lista[i].strip() if i < len(lista) and lista[i].strip() else None

    for i, nombre_area in enumerate(area_formacion_l):
        if not nombre_area.strip():
            continue  # bloque de área vacío (ej. clonado y no rellenado)

        area = ProyectoArea(
            proyecto_id=proyecto.id,
            tema_indispensable=campo(tema_indispensable_l, i),
            area_formacion=_capitalizar_corto(nombre_area.strip()),
            enfasis=campo(enfasis_l, i),
            componente=_capitalizar_corto(campo(componente_l, i)),
            contenidos=campo(contenidos_l, i),
            estrategia_pedagogica=campo(estrategia_pedagogica_l, i),
            actividades=campo(actividades_l, i),
            aprendizajes_esperados=campo(aprendizajes_esperados_l, i),
            observaciones=campo(observaciones_l, i),
        )
        db.session.add(area)
        db.session.flush()  # asigna area.id para la evaluación anidada

        evaluacion = ProyectoEvaluacion(
            area_id=area.id,
            estrategias_evaluacion=campo(estrategias_evaluacion_l, i),
            indicadores_evaluar=campo(indicadores_evaluar_l, i),
            tecnicas_evaluacion=campo(tecnicas_evaluacion_l, i),
            instrumentos_evaluacion=campo(instrumentos_evaluacion_l, i),
            tipos_evaluacion=campo(tipos_evaluacion_l, i),
            formas_participacion=campo(formas_participacion_l, i),
            recursos_materiales=campo(recursos_materiales_l, i),
        )
        db.session.add(evaluacion)


def _areas_data_json(proyecto):
    """
    Serializa las Áreas de Formación (y su Evaluación anidada) de `proyecto`
    a una lista de dicts simple, lista para pasarla al template y que el JS
    del formulario pre-cargue un bloque de área por cada una -en modo
    edición-. Devuelve lista vacía si `proyecto` es None (creación).
    """
    if not proyecto:
        return []

    datos = []
    for area in proyecto.areas:
        evaluacion = area.evaluaciones[0] if area.evaluaciones else None
        datos.append({
            'tema_indispensable': area.tema_indispensable or '',
            'area_formacion': area.area_formacion or '',
            'enfasis': area.enfasis or '',
            'componente': area.componente or '',
            'contenidos': area.contenidos or '',
            'estrategia_pedagogica': area.estrategia_pedagogica or '',
            'actividades': area.actividades or '',
            'aprendizajes_esperados': area.aprendizajes_esperados or '',
            'observaciones': area.observaciones or '',
            'estrategias_evaluacion': evaluacion.estrategias_evaluacion if evaluacion else '',
            'indicadores_evaluar': evaluacion.indicadores_evaluar if evaluacion else '',
            'tecnicas_evaluacion': evaluacion.tecnicas_evaluacion if evaluacion else '',
            'instrumentos_evaluacion': evaluacion.instrumentos_evaluacion if evaluacion else '',
            'tipos_evaluacion': evaluacion.tipos_evaluacion if evaluacion else '',
            'formas_participacion': evaluacion.formas_participacion if evaluacion else '',
            'recursos_materiales': evaluacion.recursos_materiales if evaluacion else '',
        })
    return datos


@academico_bp.route('/proyecto_aprendizaje/nuevo', methods=['GET', 'POST'])
def crear_proyecto_aprendizaje():
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    if request.method == 'POST':
        grado_id = request.form.get('grado_id')
        tema = request.form.get('tema', '').strip()
        momento_pedagogico = request.form.get('momento_pedagogico')

        if not grado_id or not tema or not momento_pedagogico:
            flash('Debe indicar el grado, el tema y el momento pedagógico.', 'danger')
            return redirect(url_for('academico.crear_proyecto_aprendizaje'))

        proyecto = ProyectoAprendizaje(docente_id=session['usuario_id'])
        _volcar_datos_generales_proyecto(proyecto)
        db.session.add(proyecto)
        db.session.flush()  # asigna proyecto.id sin cerrar la transacción

        _reconstruir_areas_proyecto(proyecto)

        db.session.commit()
        flash('Proyecto de Aprendizaje creado correctamente.', 'success')
        return redirect(url_for('academico.mi_aula', grado_id=grado_id))

    grados = Grado.query.order_by(Grado.nombre).all()
    return render_template('academico/crear_proyecto_aprendizaje.html', grados=grados, proyecto=None, areas_data=[])


@academico_bp.route('/proyecto_aprendizaje/editar/<int:id>', methods=['GET', 'POST'])
def editar_proyecto_aprendizaje(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    proyecto = ProyectoAprendizaje.query.get_or_404(id)

    es_propietario = proyecto.docente_id == session.get('usuario_id')
    es_admin = session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)']
    if not (es_propietario or es_admin):
        flash('No tiene permiso para editar este proyecto.', 'danger')
        return redirect(url_for('academico.mi_aula'))

    if request.method == 'POST':
        grado_id = request.form.get('grado_id')
        tema = request.form.get('tema', '').strip()
        momento_pedagogico = request.form.get('momento_pedagogico')

        if not grado_id or not tema or not momento_pedagogico:
            flash('Debe indicar el grado, el tema y el momento pedagógico.', 'danger')
            return redirect(url_for('academico.editar_proyecto_aprendizaje', id=id))

        _volcar_datos_generales_proyecto(proyecto)
        _reconstruir_areas_proyecto(proyecto)

        db.session.commit()
        flash('Proyecto de Aprendizaje actualizado correctamente.', 'success')
        return redirect(url_for('academico.mi_aula', grado_id=grado_id))

    grados = Grado.query.order_by(Grado.nombre).all()
    return render_template('academico/crear_proyecto_aprendizaje.html', grados=grados, proyecto=proyecto, areas_data=_areas_data_json(proyecto))


@academico_bp.route('/proyecto_aprendizaje/entregar/<int:id>', methods=['POST'])
def entregar_proyecto_aprendizaje(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    proyecto = ProyectoAprendizaje.query.get_or_404(id)

    es_propietario = proyecto.docente_id == session.get('usuario_id')
    es_admin = session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)']
    if not (es_propietario or es_admin):
        flash('No tiene permiso para marcar este proyecto como entregado.', 'danger')
        return redirect(url_for('academico.mi_aula'))

    proyecto.entregado = True
    proyecto.fecha_entrega_coordinacion = datetime.utcnow()
    db.session.commit()
    flash('Proyecto marcado como entregado a Coordinación.', 'success')
    return redirect(url_for('academico.mi_aula'))


@academico_bp.route('/coordinacion_pedagogica')
def coordinacion_pedagogica():
    """
    Panel de Coordinación Pedagógica: lista de solo lectura de los Proyectos
    de Aprendizaje ya entregados por los docentes -más sus métricas base-,
    para que Coordinación y el equipo administrativo de apoyo puedan
    revisarlos y descargar su PDF sin entrar al aula de cada uno.
    """
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    # Roles con acceso: el rango directivo/administrativo del plantel más el
    # rol dedicado de Coordinación Pedagógica (aún no dado de alta como
    # usuario en este entorno, pero ya validado aquí para cuando se cree).
    rol = session.get('nombre_rol')
    if rol not in ['Administrador Supremo', 'Equipo Directivo (Dirección)', 'Coordinador Pedagógico', 'Administrativo']:
        flash('Acceso denegado.', 'danger')
        return redirect(url_for('academico.mi_aula'))

    proyectos_entregados = ProyectoAprendizaje.query.filter_by(entregado=True) \
        .order_by(ProyectoAprendizaje.fecha_entrega_coordinacion.desc()).all()

    total_proyectos = ProyectoAprendizaje.query.filter_by(entregado=True).count()
    total_matricula = Estudiante.query.count()

    return render_template(
        'academico/coordinacion.html',
        proyectos_entregados=proyectos_entregados,
        total_proyectos=total_proyectos,
        total_matricula=total_matricula,
    )


@academico_bp.route('/generar_proyecto_prueba')
def generar_proyecto_prueba():
    """
    RUTA TEMPORAL (data seeder): inserta un Proyecto de Aprendizaje completo
    con datos realistas -tema, narrativa larga, un área y su plan de
    evaluación- para poder revisar el PDF final sin tener que llenar el
    formulario a mano. Se asigna al docente de la sesión activa.
    """
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    usuario_id = session['usuario_id']
    grado = Grado.query.filter(Grado.docentes.any(id=usuario_id)).first() or Grado.query.first()
    if not grado:
        flash('No hay ningún grado registrado en el sistema para asignar el proyecto de prueba.', 'danger')
        return redirect(url_for('academico.mi_aula'))

    lorem_diagnostico = (
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. "
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. "
        "Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum. "
        "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, totam rem aperiam."
    )
    lorem_proposito = (
        "Curabitur pretium tincidunt lacus, et sagittis magna porttitor eget. Aliquam erat volutpat. "
        "Nunc eget lorem dolor sed viverra ipsum nunc aliquet bibendum enim facilisis gravida neque. "
        "Cras ultricies ligula sed magna dictum porta, morbi leo risus, porta ac consectetur ac vestibulum. "
        "Vivamus suscipit tortor eget felis porttitor volutpat, curabitur aliquet quam id dui posuere blandit. "
        "Praesent sapien massa, convallis a pellentesque nec, egestas non nisi, id cursus nisi."
    )

    proyecto = ProyectoAprendizaje(
        docente_id=usuario_id,
        grado_id=grado.id,
        tema='SISTEMA DEL CUERPO HUMANO',
        momento_pedagogico='I',
        fecha_inicio=date.today(),
        fecha_culminacion=date.today() + timedelta(days=30),
        fecha_entrega=date.today() + timedelta(days=35),
        diagnostico_pedagogico=lorem_diagnostico,
        proposito_integral=lorem_proposito,
        cierre_demostracion='Los estudiantes realizarán una demostración práctica sobre el funcionamiento del sistema muscular y óseo del cuerpo humano.',
        cierre_dramatizacion='Dramatización de una consulta médica donde se explican, en lenguaje sencillo, los sistemas del cuerpo humano trabajados.',
        cierre_muestra='Exposición de maquetas y carteleras elaboradas por los estudiantes sobre el sistema del cuerpo humano.',
    )
    db.session.add(proyecto)
    db.session.flush()  # asigna proyecto.id sin cerrar la transacción

    # Contenidos y sus campos "hermanos" (Estrategia/Actividades/Aprendizajes)
    # se cargan con 3 ítems sincronizados uno a uno -así se puede comprobar
    # que la lógica de "compartimentos" del PDF funciona: cada ítem debe
    # imprimirse en su propia fila apilada, con línea divisoria propia.
    area = ProyectoArea(
        proyecto_id=proyecto.id,
        tema_indispensable='PRESERVACIÓN DE LA VIDA EN EL PLANETA',
        area_formacion='CIENCIAS NATURALES',
        enfasis='Educación ambiental y promoción de la salud integral',
        componente='*Salud',
        contenidos='1. Sistema Muscular\n2. Sistema Óseo\n3. Sistema Nervioso',
        estrategia_pedagogica='1. Taller en grupo\n2. Demostración práctica\n3. Juego didáctico',
        actividades=(
            '1. El Esqueleto de Cartón: Construir un brazo mecánico utilizando cartón, sujetadores mariposa y cordones.\n'
            '2. Palpación guiada de los huesos principales del cuerpo sobre una lámina anatómica.\n'
            '3. Dinámica "Simón dice" moviendo distintos grupos musculares al nombrarlos.'
        ),
        aprendizajes_esperados=(
            '1. Comprende que el movimiento del cuerpo depende de la acción conjunta de huesos y músculos.\n'
            '2. Identifica los huesos principales del esqueleto humano.\n'
            '3. Reconoce la importancia de cuidar el sistema nervioso.'
        ),
        observaciones='',
    )
    db.session.add(area)
    db.session.flush()  # asigna area.id para la evaluación anidada

    evaluacion = ProyectoEvaluacion(
        area_id=area.id,
        estrategias_evaluacion=(
            '1. Observación directa durante la construcción del brazo mecánico.\n'
            '2. Participación activa en el taller grupal.\n'
            '3. Autoevaluación oral al cierre de la actividad.'
        ),
        indicadores_evaluar=(
            '1. Identifica los huesos y músculos principales que intervienen en el movimiento del brazo.\n'
            '2. Explica con sus palabras cómo se mueve el cuerpo humano.\n'
            '3. Trabaja en equipo respetando el turno de sus compañeros.'
        ),
        tecnicas_evaluacion='Observación y análisis de producto',
        instrumentos_evaluacion='Lista de cotejo',
        tipos_evaluacion='Formativa',
        formas_participacion='Coevaluación',
        recursos_materiales='Cartón, sujetadores mariposa, cordones, tijeras, marcadores',
    )
    db.session.add(evaluacion)

    db.session.commit()
    flash('Proyecto de prueba "SISTEMA DEL CUERPO HUMANO" creado correctamente.', 'success')
    return redirect(url_for('academico.mi_aula'))


def _dividir_texto_pdf(pdf, texto, ancho_util):
    """
    Envuelve `texto` en líneas que quepan dentro de `ancho_util` mm, replicando
    a mano el wrap de multi_cell (PyFPDF 1.7.2 no expone un modo "dry run" para
    medir cuántas líneas ocupará un texto sin imprimirlo). Se usa para calcular
    la altura de cada celda ANTES de dibujar la fila de la tabla, de modo que
    todas las columnas de una misma fila queden con el mismo alto.
    """
    texto = (texto or '').strip()
    if not texto:
        return ['-']

    lineas = []
    for parrafo in texto.split('\n'):
        parrafo = parrafo.strip()
        if not parrafo:
            lineas.append('')
            continue
        actual = ''
        for palabra in parrafo.split(' '):
            candidato = f"{actual} {palabra}".strip()
            if pdf.get_string_width(candidato) <= ancho_util:
                actual = candidato
                continue
            if actual:
                lineas.append(actual)
            if pdf.get_string_width(palabra) <= ancho_util:
                actual = palabra
            else:
                # Palabra más ancha que la columna completa: se corta por caracteres.
                trozo = ''
                for ch in palabra:
                    if pdf.get_string_width(trozo + ch) <= ancho_util:
                        trozo += ch
                    else:
                        lineas.append(trozo)
                        trozo = ch
                actual = trozo
        if actual:
            lineas.append(actual)
    return lineas or ['-']


def _fila_tabla_pdf(pdf, x_inicio, anchos, textos, line_h=5, font=('Arial', '', 7.5), alto_min=9, borde_inferior=True):
    """
    Dibuja una fila de tabla "tipo multi_cell" con bordes regulares (sin
    rellenos de color): cada columna puede ocupar varias líneas, pero todas
    las celdas de la fila quedan con la misma altura (la de la columna con
    más texto). Las cabeceras de tabla se distinguen solo por fuente en
    negrita, pasando `font=('Arial', 'B', N)`. `alto_min` garantiza que cada
    fila luzca como un "cuadro predeterminado" (altura generosa) aunque el
    texto sea corto. `borde_inferior=False` deja las columnas "abiertas"
    hacia abajo (sin la raya horizontal justo debajo del texto) -para filas
    que luego se prolongan con `_anclar_fondo_tabla_pdf` y no deben mostrar
    un cierre duplicado-. Hace salto de página automático si la fila no cabe
    en el espacio restante. Devuelve el alto (mm) dibujado.
    """
    pdf.set_font(*font)
    lineas_por_columna = [_dividir_texto_pdf(pdf, texto, ancho - 2) for texto, ancho in zip(textos, anchos)]
    num_lineas = max(len(lineas) for lineas in lineas_por_columna)
    alto_fila = max(num_lineas * line_h + 1.5, alto_min)

    if pdf.get_y() + alto_fila > pdf.h - pdf.b_margin:
        pdf.add_page()

    y = pdf.get_y()
    x = x_inicio
    pdf.set_draw_color(120, 120, 120)
    pdf.set_text_color(0, 0, 0)
    for ancho, lineas in zip(anchos, lineas_por_columna):
        if borde_inferior:
            pdf.rect(x, y, ancho, alto_fila)
        else:
            pdf.line(x, y, x, y + alto_fila)                    # borde izquierdo
            pdf.line(x + ancho, y, x + ancho, y + alto_fila)    # borde derecho
            pdf.line(x, y, x + ancho, y)                        # borde superior (sin el inferior)
        y_linea = y + 1.2
        for linea in lineas:
            pdf.set_xy(x + 1, y_linea)
            pdf.cell(ancho - 2, line_h, linea, border=0, align='L')
            y_linea += line_h
        x += ancho

    pdf.set_xy(x_inicio, y + alto_fila)
    return alto_fila


def _anclar_fondo_tabla_pdf(pdf, x_inicio, anchos, y_fondo):
    """
    "Cuadrícula anclada": toma la Y donde terminó de imprimirse el texto de
    la última fila dibujada (pdf.get_y() en ese momento) y extiende las
    líneas verticales de cada columna -calculadas a partir de `anchos`-
    desde ahí hasta la coordenada FIJA `y_fondo`, cerrando con una línea
    horizontal en esa misma Y. Así la tabla siempre llega hasta el fondo de
    la hoja (o hasta la mitad, según `y_fondo`), sin importar si el docente
    escribió poco texto. No hace nada si el contenido real ya alcanzó o
    superó `y_fondo` (no se dibujan líneas de longitud negativa).
    """
    current_y = pdf.get_y()
    if y_fondo <= current_y:
        return

    pdf.set_draw_color(0, 0, 0)
    x = x_inicio
    pdf.line(x, current_y, x, y_fondo)
    for ancho in anchos:
        x += ancho
        pdf.line(x, current_y, x, y_fondo)

    pdf.line(x_inicio, y_fondo, x, y_fondo)
    pdf.set_xy(x_inicio, y_fondo)


def _limites_columnas_pdf(x_inicio, anchos):
    """Devuelve la lista de coordenadas X acumuladas (bordes de cada columna),
    empezando en `x_inicio` y terminando en el borde derecho del bloque."""
    limites = [x_inicio]
    for ancho in anchos:
        limites.append(limites[-1] + ancho)
    return limites


def _dividir_items_pdf(texto):
    """
    Divide un campo de texto en una lista de "ítems" (uno por línea no
    vacía), para simular varias filas/compartimentos dentro de una misma
    columna -ej. una lista numerada "1. ...\\n2. ...\\n3. ...". Si el texto
    no tiene saltos de línea reales, se trata como un único ítem.
    """
    items = [linea.strip() for linea in (texto or '').split('\n') if linea.strip()]
    return items or ['-']


def _imprimir_valor_columna_pdf(pdf, x, ancho, texto, y_inicio, font, line_h, y_fondo=None):
    """
    Imprime UN SOLO valor (puede ocupar varias líneas por word-wrap) en una
    columna, sin bordes propios -para las columnas que son "una sola por
    página" (ej. Área/Énfasis/Componente, o Técnicas/Instrumentos)-.

    Si `y_fondo` se indica, el valor se centra verticalmente dentro del
    compartimento fijo que va desde `y_inicio` hasta `y_fondo` (altura
    disponible menos la altura real que ocupa el texto, repartida como
    margen superior). Si se omite, el texto queda alineado arriba (comportamiento
    original).
    """
    pdf.set_font(*font)
    lineas = _dividir_texto_pdf(pdf, texto or '-', ancho - 2)

    if y_fondo is not None:
        altura_disponible = y_fondo - y_inicio
        altura_texto = len(lineas) * line_h
        margen_y = max((altura_disponible - altura_texto) / 2, 0)
        y_linea = y_inicio + margen_y
    else:
        y_linea = y_inicio + 1.5

    for linea in lineas:
        pdf.set_xy(x + 1, y_linea)
        pdf.cell(ancho - 2, line_h, linea, border=0, align='L')
        y_linea += line_h


def _dibujar_compartimentos_pdf(pdf, y_inicio, x_divisor_izq, x_divisor_der, columnas, font, line_h):
    """
    Dibuja los "compartimentos" de un grupo de columnas que comparten el
    mismo listado de ítems paralelos (ej. Contenidos/Estrategia/Actividades/
    Aprendizajes/Observaciones, o Estrategias/Indicadores). `columnas` es una
    lista de tuplas (x, ancho, items); se usa la lista más larga como
    "maestra" para el número de filas. Después de cada fila se dibuja una
    línea horizontal OBLIGATORIA entre `x_divisor_izq` y `x_divisor_der`
    -y NUNCA más allá de ese rango-, creando el efecto visual de
    compartimentos apilados. Devuelve la Y donde terminó el último.
    """
    pdf.set_font(*font)
    num_filas = max(len(items) for _, _, items in columnas)
    y_cursor = y_inicio
    for i in range(num_filas):
        y_fila_inicio = y_cursor
        y_fin_fila = y_fila_inicio
        for x, ancho, items in columnas:
            texto = items[i] if i < len(items) else ''
            if not texto:
                continue
            lineas = _dividir_texto_pdf(pdf, texto, ancho - 2)
            y_linea = y_fila_inicio + 1.5
            for linea in lineas:
                pdf.set_xy(x + 1, y_linea)
                pdf.cell(ancho - 2, line_h, linea, border=0, align='L')
                y_linea += line_h
            y_fin_fila = max(y_fin_fila, y_linea)

        if y_fin_fila == y_fila_inicio:
            y_fin_fila += line_h  # compartimento vacío: igual reserva una línea de alto

        pdf.set_draw_color(0, 0, 0)
        pdf.line(x_divisor_izq, y_fin_fila, x_divisor_der, y_fin_fila)
        y_cursor = y_fin_fila

    return y_cursor


def _bloque_metadatos_portada_pdf(pdf, margin, filas, x_izq=15, x_der=160, col_w=105, alto_fila=8, font_size=12, y_offset=10):
    """
    Bloque de metadatos de la portada: columnas separadas radicalmente hacia
    los márgenes de la hoja (NO centradas en un bloque angosto) -columna
    izquierda forzada en X=`x_izq`, columna derecha forzada en X=`x_der`-.
    Cada fila es una tupla (par_izquierdo, par_derecho); `par_derecho` puede
    ser None para dejar esa mitad vacía (ej. la fila del Docente). Cada par
    es (etiqueta, valor). `y_offset` baja el punto de partida del bloque
    para separarlo del título que va justo encima.
    """
    pdf.ln(y_offset)
    pdf.set_text_color(0, 0, 0)
    for izquierda, derecha in filas:
        y_fila = pdf.get_y()
        for x, par in ((x_izq, izquierda), (x_der, derecha)):
            if par is None:
                continue
            etiqueta, valor = par
            pdf.set_xy(x, y_fila)
            pdf.set_x(x)  # fuerza explícitamente el arranque de la columna en X
            pdf.set_font('Arial', 'B', font_size)
            ancho_etiqueta = pdf.get_string_width(etiqueta) + 2
            pdf.cell(ancho_etiqueta, alto_fila, etiqueta, ln=0)
            pdf.set_font('Arial', '', font_size)
            pdf.cell(col_w - ancho_etiqueta, alto_fila, valor, ln=0)
        pdf.set_xy(margin, y_fila + alto_fila)


def _dibujar_cabecera_proyecto_pdf(pdf, margin):
    """
    Membrete institucional de la portada: logos a los costados y debajo la
    República, el Ministerio, el nombre del plantel y la ciudad, todos
    centrados. NO imprime el título del documento -eso lo hace el caller
    después de su propio salto de línea, para controlar el espaciado
    milimétrico de la portada-. Devuelve la Y donde terminó de imprimirse.
    """
    page_w = pdf.w
    usable_w = page_w - (2 * margin)

    logo_ministerio_w, logo_ministerio_h = 65, 26
    logo_eep_w, logo_eep_h = 40, 40
    logo_h_max = max(logo_ministerio_h, logo_eep_h)
    logo_inset = 6
    logo_ministerio = os.path.join(current_app.root_path, 'static', 'img', 'LogoMInisterioDeEducacion.png')
    logo_eep = os.path.join(current_app.root_path, 'static', 'img', 'LogoEEP.png')
    y_logo = margin
    y_logo_ministerio = y_logo + (logo_h_max - logo_ministerio_h) / 2
    y_logo_eep = y_logo + (logo_h_max - logo_eep_h) / 2
    logo_x_izq = margin + logo_inset
    logo_x_der = page_w - margin - logo_inset - logo_eep_w

    if os.path.exists(logo_ministerio):
        pdf.image(logo_ministerio, x=logo_x_izq, y=y_logo_ministerio, w=logo_ministerio_w, h=logo_ministerio_h, type=_tipo_imagen(logo_ministerio))
    if os.path.exists(logo_eep):
        pdf.image(logo_eep, x=logo_x_der, y=y_logo_eep, w=logo_eep_w, h=logo_eep_h, type=_tipo_imagen(logo_eep))

    # Todo el membrete al mismo tamaño (14pt), incluida la ciudad -en negro,
    # como el resto-, para que se lea como un bloque institucional uniforme.
    pdf.set_xy(margin, y_logo)
    pdf.set_font('Arial', 'B', 14)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(usable_w, 6.5, 'República Bolivariana de Venezuela', ln=2, align='C')
    pdf.set_x(margin)
    pdf.cell(usable_w, 6.5, 'Ministerio del Poder Popular para la Educación', ln=2, align='C')
    pdf.set_x(margin)
    pdf.cell(usable_w, 6.5, "Escuela de Educación Primaria Daniel O'Leary", ln=2, align='C')
    pdf.set_x(margin)
    pdf.set_font('Arial', '', 14)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(usable_w, 6.5, 'San Fernando-Edo Apure', ln=2, align='C')

    pdf.set_xy(margin, y_logo + logo_h_max)
    return pdf.get_y()


def _construir_pdf_proyecto_aprendizaje(proyecto):
    """
    Genera el PDF oficial (Carta, horizontal) del Proyecto de Aprendizaje.
    Paginación estricta: portada sola, narrativa sola, y por cada área DOS
    hojas (una para la matriz de Áreas de Formación, otra -obligatoria- para
    su Plan de Evaluación apilado), y cierre solo en la última.
    """
    margin = 10
    pdf = FPDF(orientation='L', unit='mm', format='Letter')
    pdf.set_auto_page_break(auto=True, margin=margin)
    pdf.set_margins(margin, margin, margin)
    usable_w = pdf.w - (2 * margin)

    # Margen lateral compartido por la portada (bloque de metadatos) y la
    # narrativa ('Diagnóstico Pedagógico' / 'Propósito Integral'), para que
    # ambas páginas queden alineadas exactamente en el mismo margen X -una
    # sola fuente de verdad en vez de dos números repetidos por separado-.
    margin_narrativa = margin + 10
    ancho_narrativa = usable_w - 2 * 10

    # ── Matrícula dinámica (estudiantes activos del grado del proyecto) ──
    estudiantes_activos = Estudiante.query.filter_by(grado_id=proyecto.grado_id, estatus='Activo').all()
    hembras = sum(1 for e in estudiantes_activos if e.genero == 'Femenino')
    varones = sum(1 for e in estudiantes_activos if e.genero == 'Masculino')
    total_matricula = len(estudiantes_activos)

    docente_nombre = proyecto.docente.nombre_completo if proyecto.docente else 'No asignado'
    grado_nombre = proyecto.grado.nombre if proyecto.grado else '-'
    fecha_inicio = proyecto.fecha_inicio.strftime('%d/%m/%Y') if proyecto.fecha_inicio else '--/--/----'
    fecha_culminacion = proyecto.fecha_culminacion.strftime('%d/%m/%Y') if proyecto.fecha_culminacion else '--/--/----'

    # ══════════════════════════ PÁGINA 1: PORTADA (exclusiva) ══════════════════════════
    pdf.add_page()
    pdf.set_font('Arial', '', 12)
    _dibujar_cabecera_proyecto_pdf(pdf, margin)

    # Título dinámico: si el proyecto tiene tema, se imprime SOLO el tema (en
    # mayúsculas); el genérico 'PROYECTO DE APRENDIZAJE' solo aparece como
    # respaldo cuando no hay tema cargado.
    pdf.set_y(70)
    pdf.set_x(margin)
    pdf.set_font('Arial', 'B', 20)
    pdf.set_text_color(0, 0, 0)
    if proyecto.tema:
        pdf.cell(usable_w, 10, proyecto.tema.upper(), ln=1, align='C')
    else:
        pdf.cell(usable_w, 10, 'PROYECTO DE APRENDIZAJE', ln=1, align='C')

    # Bloque de metadatos, empujado al centro-abajo de la hoja (Y=120),
    # independiente de dónde terminó el título, para que la portada quede
    # repartida en cuartos: Membrete / Título / Metadatos / Firmas.
    #
    # Márgenes laterales: la columna izquierda arranca exactamente en
    # `margin_narrativa` -el mismo margen X que usa 'Diagnóstico Pedagógico'
    # en la página 2-, y la columna derecha es su espejo exacto respecto al
    # borde de la hoja (pdf.w - margin_narrativa - col_w), de modo que el
    # bloque completo quede tan simétrico y ancho como el de la narrativa.
    metadatos_col_w = 100
    x_izq_metadatos = margin_narrativa
    x_der_metadatos = pdf.w - margin_narrativa - metadatos_col_w
    pdf.set_y(120)
    _bloque_metadatos_portada_pdf(pdf, margin, [
        (('GRADO: ', grado_nombre), ('MOMENTO PEDAGÓGICO: ', proyecto.momento_pedagogico or '-')),
        (('MATRÍCULA: ', f"H: {hembras}   V: {varones}   TOTAL: {total_matricula}"), ('FECHA DE CULMINACIÓN: ', fecha_culminacion)),
        (('FECHA DE INICIO: ', fecha_inicio), ('DOCENTE: ', docente_nombre)),
    ], x_izq=x_izq_metadatos, x_der=x_der_metadatos, col_w=metadatos_col_w, font_size=12, y_offset=0)

    # Firmas ancladas al fondo de la hoja (posición absoluta, no relativa a
    # dónde terminó el bloque de metadatos), para que siempre queden abajo.
    pdf.set_y(175)
    pdf.set_x(margin)
    pdf.set_font('Arial', '', 10)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(80, 10, 'ENLACE PEDAGÓGICO: ____________', align='C')
    pdf.cell(80, 10, 'DIRECTORA (E) DAISSY REYNA: ____________', align='C')
    pdf.cell(80, 10, 'FECHA DE ENTREGA: ____________', align='C')

    # ══════════════════════════ PÁGINA 2: NARRATIVA (exclusiva) ══════════════════════════
    pdf.add_page()

    # 'Diagnóstico Pedagógico' sube, más cerca del margen superior.
    # (margin_narrativa / ancho_narrativa se definieron al inicio de la
    # función para compartirlos con el bloque de metadatos de la portada).
    pdf.set_y(25)
    pdf.set_x(margin_narrativa)
    pdf.set_font('Arial', 'B', 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(ancho_narrativa, 7, 'Diagnóstico Pedagógico', ln=1, align='C')
    pdf.set_x(margin_narrativa)
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(ancho_narrativa, 5, proyecto.diagnostico_pedagogico or 'No especificado.', align='J')

    # 'Propósito Integral' arranca bien separado del bloque anterior (buen
    # espaciado vertical entre ambos), en Y=95.
    pdf.set_y(95)
    pdf.set_x(margin_narrativa)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(ancho_narrativa, 7, 'Propósito Integral', ln=1, align='C')
    pdf.set_x(margin_narrativa)
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(ancho_narrativa, 5, proyecto.proposito_integral or 'No especificado.', align='J')

    # ══════════════════════ CICLO DE ÁREAS: una página nueva por área ══════════════════════
    # COMPONENTE +4 puntos (9->13) para que la palabra quepa en una sola
    # línea de cabecera; se le restan a OBSERVACIONES (16->12).
    anchos_area = [p / 100 * usable_w for p in (10, 9, 13, 14, 14, 14, 14, 12)]
    encabezados_area = [
        'ÁREAS DE FORMACIÓN', 'ÉNFASIS', 'COMPONENTE', 'CONTENIDOS',
        'ESTRATEGIA PEDAGÓGICA', 'ACTIVIDADES', 'APRENDIZAJES ESPERADOS', 'OBSERVACIONES',
    ]

    anchos_eval_sup = [usable_w / 4] * 4
    encabezados_eval_sup = ['ESTRATEGIAS', 'INDICADORES', 'TÉCNICAS', 'INSTRUMENTOS']

    anchos_eval_inf = [usable_w / 3] * 3
    encabezados_eval_inf = ['TIPOS', 'FORMAS DE PARTICIPACIÓN', 'RECURSOS']

    if not proyecto.areas:
        pdf.add_page()
        pdf.set_x(margin)
        pdf.set_font('Arial', 'I', 10)
        pdf.cell(usable_w, 6, 'Este proyecto no tiene áreas de formación registradas.', ln=1)

    # Fuente fija para las celdas del Plan de Evaluación (no lleva auto-escalado).
    FONT_EVAL_DATO = ('Arial', '', 12)
    LINE_H_EVAL = 6

    limites_area = _limites_columnas_pdf(margin, anchos_area)
    x_contenidos_inicio = limites_area[3]   # borde izquierdo de la columna CONTENIDOS
    x_borde_derecho_area = limites_area[-1]  # borde derecho de la última columna (OBSERVACIONES)

    limites_eval_sup = _limites_columnas_pdf(margin, anchos_eval_sup)
    x_estrategias_inicio = limites_eval_sup[0]
    x_indicadores_fin = limites_eval_sup[2]  # borde derecho de INDICADORES (antes de TÉCNICAS)

    for area in proyecto.areas:
        pdf.add_page()

        # "TEMAS INDISPENSABLES" ya no es texto rojo suelto: es una celda con
        # borde que ocupa todo el ancho de la tabla y hace de "techo" de la
        # matriz -las cabeceras se dibujan inmediatamente debajo, sin saltos-.
        pdf.set_x(margin)
        pdf.set_font('Arial', 'B', 10)
        pdf.set_text_color(0, 0, 0)
        tema_indisp = area.tema_indispensable or 'No especificado'
        pdf.cell(usable_w, 7, f"TEMAS INDISPENSABLES: {tema_indisp.upper()}", border=1, ln=1, align='L')

        # Cabecera de la matriz (fuente más grande, 11pt) -llena mejor el espacio-.
        pdf.set_x(margin)
        _fila_tabla_pdf(pdf, margin, anchos_area, encabezados_area, font=('Arial', 'B', 11), alto_min=10)

        # ── COMPARTIMENTOS: Área/Énfasis/Componente son UN SOLO valor por
        # página; Contenidos y sus campos hermanos pueden ser varios ítems
        # (uno por línea) y se dibujan como filas apiladas, cada una cerrada
        # por una línea horizontal que va SOLO de Contenidos al borde derecho
        # -nunca cruza sobre Área/Énfasis/Componente-. ──
        y_inicio_datos = pdf.get_y()

        # Auto-escalado de fuente según cuántos contenidos hay, para evitar
        # desbordamientos cuando el docente carga muchos ítems.
        contenidos_items = _dividir_items_pdf(area.contenidos)
        if len(contenidos_items) <= 3:
            font_area_datos = ('Arial', '', 10)
            line_h_area = 5
        else:
            font_area_datos = ('Arial', '', 8)
            line_h_area = 4.2

        _imprimir_valor_columna_pdf(pdf, limites_area[0], anchos_area[0], area.area_formacion, y_inicio_datos, font_area_datos, line_h_area, y_fondo=185)
        _imprimir_valor_columna_pdf(pdf, limites_area[1], anchos_area[1], area.enfasis, y_inicio_datos, font_area_datos, line_h_area, y_fondo=185)
        _imprimir_valor_columna_pdf(pdf, limites_area[2], anchos_area[2], area.componente, y_inicio_datos, font_area_datos, line_h_area, y_fondo=185)

        columnas_compartimento = [
            (limites_area[3], anchos_area[3], contenidos_items),
            (limites_area[4], anchos_area[4], _dividir_items_pdf(area.estrategia_pedagogica)),
            (limites_area[5], anchos_area[5], _dividir_items_pdf(area.actividades)),
            (limites_area[6], anchos_area[6], _dividir_items_pdf(area.aprendizajes_esperados)),
            (limites_area[7], anchos_area[7], _dividir_items_pdf(area.observaciones)),
        ]
        _dibujar_compartimentos_pdf(pdf, y_inicio_datos, x_contenidos_inicio, x_borde_derecho_area, columnas_compartimento, font_area_datos, line_h_area)

        # Cuadrícula anclada: las verticales de TODAS las columnas se dibujan
        # desde el tope de la fila (y_inicio_datos) hasta el fondo fijo
        # Y=185, cerrando con una horizontal completa -así Área/Énfasis/
        # Componente se ven como una única celda alta, y toda la matriz llena
        # la página sin importar cuánto texto real haya-.
        pdf.set_y(y_inicio_datos)
        _anclar_fondo_tabla_pdf(pdf, margin, anchos_area, y_fondo=185)

        # Salto de página OBLIGATORIO: el Plan de Evaluación de esta área
        # siempre arranca en una hoja nueva, separada de su matriz. Cada área
        # ocupa entonces 2 hojas completas.
        pdf.add_page()

        pdf.set_x(margin)
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(usable_w, 6, 'PLAN DE EVALUACIÓN', ln=1, align='C')

        evaluacion = area.evaluaciones[0] if area.evaluaciones else None

        if evaluacion:
            # ── Tabla superior: Estrategias / Indicadores / Técnicas / Instrumentos.
            # Técnicas e Instrumentos son UN SOLO valor; Estrategias e Indicadores
            # se compartimentan igual que la matriz (línea divisoria SOLO entre
            # Estrategias e Indicadores, sin invadir Técnicas/Instrumentos).
            # RE-PROPORCIÓN: esta tabla ahora queda anclada hasta Y=150, dándole
            # mucho más espacio (antes la inferior le "robaba" espacio de sobra).
            pdf.set_x(margin)
            _fila_tabla_pdf(pdf, margin, anchos_eval_sup, encabezados_eval_sup, font=('Arial', 'B', 11), alto_min=10)

            y_inicio_sup = pdf.get_y()
            _imprimir_valor_columna_pdf(pdf, limites_eval_sup[2], anchos_eval_sup[2], evaluacion.tecnicas_evaluacion, y_inicio_sup, FONT_EVAL_DATO, LINE_H_EVAL)
            _imprimir_valor_columna_pdf(pdf, limites_eval_sup[3], anchos_eval_sup[3], evaluacion.instrumentos_evaluacion, y_inicio_sup, FONT_EVAL_DATO, LINE_H_EVAL)

            columnas_compartimento_eval = [
                (limites_eval_sup[0], anchos_eval_sup[0], _dividir_items_pdf(evaluacion.estrategias_evaluacion)),
                (limites_eval_sup[1], anchos_eval_sup[1], _dividir_items_pdf(evaluacion.indicadores_evaluar)),
            ]
            _dibujar_compartimentos_pdf(pdf, y_inicio_sup, x_estrategias_inicio, x_indicadores_fin, columnas_compartimento_eval, FONT_EVAL_DATO, LINE_H_EVAL)

            pdf.set_y(y_inicio_sup)
            _anclar_fondo_tabla_pdf(pdf, margin, anchos_eval_sup, y_fondo=150)

            # Salto fuerte: la tabla inferior arranca en Y=155 (deja aire tras
            # el cierre en Y=150 de la tabla superior) y queda anclada hasta
            # Y=185 -un bloque pequeño, solo para el cierre de la hoja-.
            pdf.set_y(155)
            pdf.set_x(margin)
            _fila_tabla_pdf(pdf, margin, anchos_eval_inf, encabezados_eval_inf, font=('Arial', 'B', 11), alto_min=10)
            # Sin borde inferior: las columnas quedan "abiertas" hacia abajo,
            # limitadas solo por las verticales y el cierre final del anclaje
            # en Y=185 -evita la raya sobrante justo debajo del texto-.
            pdf.set_x(margin)
            _fila_tabla_pdf(pdf, margin, anchos_eval_inf, [
                evaluacion.tipos_evaluacion or '-',
                evaluacion.formas_participacion or '-',
                evaluacion.recursos_materiales or '-',
            ], font=FONT_EVAL_DATO, borde_inferior=False)

            _anclar_fondo_tabla_pdf(pdf, margin, anchos_eval_inf, y_fondo=185)
        else:
            pdf.set_x(margin)
            pdf.set_font('Arial', 'I', 9)
            pdf.cell(usable_w, 6, 'Esta área no tiene un plan de evaluación registrado.', ln=1)

    # ══════════════════════════ PÁGINA FINAL: CIERRE (exclusiva) ══════════════════════════
    pdf.add_page()
    pdf.set_x(margin)
    pdf.set_font('Arial', 'B', 14)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(usable_w, 8, 'CIERRE DEL PROYECTO', ln=1, align='C')
    pdf.ln(4)

    # Demostración / Dramatización / Muestra Cultural: tres cuadros grandes con
    # borde (no texto libre). Se rellena con líneas en blanco hasta un mínimo
    # de MIN_LINEAS para que la caja se vea "grande" aunque el texto sea corto.
    MIN_LINEAS = 7
    LINE_H = 5
    for etiqueta, texto in [
        ('Demostración', proyecto.cierre_demostracion),
        ('Dramatización', proyecto.cierre_dramatizacion),
        ('Muestra Cultural', proyecto.cierre_muestra),
    ]:
        pdf.set_x(margin)
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(usable_w, 6, etiqueta, ln=1, align='C')

        contenido = texto or 'No especificado.'
        pdf.set_font('Arial', '', 10)
        lineas = _dividir_texto_pdf(pdf, contenido, usable_w - 4)
        lineas_faltantes = max(0, MIN_LINEAS - len(lineas))
        contenido_caja = contenido + ('\n' * lineas_faltantes)

        pdf.set_x(margin)
        pdf.multi_cell(usable_w, LINE_H, contenido_caja, border=1, align='J')
        pdf.ln(4)

    return pdf


@academico_bp.route('/proyecto_aprendizaje/pdf/<int:proyecto_id>')
def generar_proyecto_aprendizaje_pdf(proyecto_id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    proyecto = ProyectoAprendizaje.query.get_or_404(proyecto_id)

    es_propietario = proyecto.docente_id == session.get('usuario_id')
    es_admin = session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)']
    if not (es_propietario or es_admin):
        flash('No tiene permiso para generar el PDF de este proyecto.', 'danger')
        return redirect(url_for('academico.mi_aula'))

    pdf = _construir_pdf_proyecto_aprendizaje(proyecto)

    grado_txt = (proyecto.grado.nombre if proyecto.grado else 'Grado').strip().replace(' ', '_')
    momento_txt = (proyecto.momento_pedagogico or '').strip().replace(' ', '_')
    nombre_archivo = f"Proyecto_Aprendizaje_{grado_txt}_{momento_txt}.pdf"

    response = make_response(pdf.output(dest='S').encode('latin1'))
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'inline; filename="{nombre_archivo}"'
    return response


@academico_bp.route('/proyecto_aprendizaje/eliminar/<int:proyecto_id>', methods=['POST'])
def eliminar_proyecto_aprendizaje(proyecto_id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    proyecto = ProyectoAprendizaje.query.get_or_404(proyecto_id)

    es_propietario = proyecto.docente_id == session.get('usuario_id')
    es_admin = session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)']
    if not (es_propietario or es_admin):
        flash('No tiene permiso para eliminar este proyecto.', 'danger')
        return redirect(url_for('academico.mi_aula'))

    db.session.delete(proyecto)  # cascada: elimina también sus áreas y evaluaciones
    db.session.commit()
    flash('Proyecto de Aprendizaje eliminado correctamente.', 'success')
    return redirect(url_for('academico.mi_aula'))


@academico_bp.route('/guardar_evaluacion/<int:estudiante_id>', methods=['POST'])
def guardar_evaluacion(estudiante_id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    est = Estudiante.query.get_or_404(estudiante_id)
    momento_val = request.form.get('momento')
    if not momento_val:
        flash('Debe indicar el momento a guardar.', 'danger')
        return redirect(request.referrer or url_for('academico.mi_aula'))

    momento = int(momento_val)
    grado_id = request.form.get('grado_id')

    evaluacion = EvaluacionEstudiante.query.filter_by(estudiante_id=est.id, momento=momento).first()
    if not evaluacion:
        evaluacion = EvaluacionEstudiante(estudiante_id=est.id, momento=momento)
        db.session.add(evaluacion)

    evaluacion.texto_descriptivo = request.form.get('texto_descriptivo', '').strip()
    evaluacion.sugerencias = request.form.get('sugerencias', '').strip()

    db.session.commit()
    flash(f'Evaluación del Momento {momento} guardada para {est.nombre_completo}.', 'success')
    return redirect(url_for('academico.mi_aula', grado_id=grado_id))


def _tipo_imagen(ruta):
    """
    fpdf 1.7.2 elige el parser por la extensión salvo que se indique `type=`
    explícitamente; algunos assets están guardados con extensión .png pero
    contenido JPEG (o viceversa), así que se detecta por la firma real de bytes.
    """
    try:
        with open(ruta, 'rb') as f:
            cabecera = f.read(8)
    except OSError:
        return ''
    if cabecera.startswith(b'\x89PNG'):
        return 'PNG'
    if cabecera.startswith(b'\xff\xd8'):
        return 'JPG'
    return ''


def _separar_grado_seccion(texto, seccion_fallback='-'):
    """
    Separa un string de grado que puede venir combinado con la sección
    (ej. "Segundo Grado Sección A") en sus dos partes: grado y sección.
    Si no se detecta la palabra "Sección" (con o sin acento), se devuelve
    el texto completo como grado y `seccion_fallback` como sección.
    """
    texto = (texto or '').strip()
    if not texto:
        return '-', seccion_fallback

    texto_normalizado = ''.join(
        c for c in unicodedata.normalize('NFD', texto) if unicodedata.category(c) != 'Mn'
    )
    match = re.search(r'secci[oó]n', texto_normalizado, flags=re.IGNORECASE)
    if not match:
        return texto, seccion_fallback

    grado_txt = texto[:match.start()].strip(' -,') or texto
    resto = texto[match.end():].strip(' -,:')
    seccion_txt = resto or seccion_fallback
    return grado_txt, seccion_txt


def _dibujar_hoja_boletin(pdf, estudiante, proyectos, evaluaciones):
    """
    Dibuja UNA hoja completa del Registro Descriptivo (Boletín) sobre la página
    actualmente activa de `pdf`. El caller debe llamar pdf.add_page() antes de
    invocar esta función (así se reutiliza igual para 1 estudiante o para un
    PDF consolidado de varias páginas, una por estudiante).

    proyectos: dict {momento: ProyectoAula} del salón del estudiante.
    evaluaciones: dict {momento: EvaluacionEstudiante} de ESTE estudiante.
    """
    est = estudiante
    grado = est.grado
    docentes_nombres = ", ".join([d.nombre_completo for d in grado.docentes]) if (grado and grado.docentes) else "Docente no asignado"

    # C.I: el estudiante no maneja cédula propia en el sistema; se fuerza siempre la Cédula Escolar (C.E.S.C).
    ci_mostrar = est.cedula_escolar or 'S/N'

    margin = 10
    page_w = pdf.w
    page_h = pdf.h
    usable_w = page_w - (2 * margin)

    # ── Logos + Membrete ──
    # El logo del Ministerio es rectangular (ancho); forzarlo a un cuadrado lo deforma.
    logo_ministerio_w = 45
    logo_ministerio_h = 18
    # El logo de la escuela sí es cuadrado.
    logo_eep_w = 25
    logo_eep_h = 25
    logo_h_max = max(logo_ministerio_h, logo_eep_h)
    logo_inset = 6  # separación extra respecto al margen lateral, para no pegar los logos al borde
    logo_ministerio = os.path.join(current_app.root_path, 'static', 'img', 'LogoMInisterioDeEducacion.png')
    logo_eep = os.path.join(current_app.root_path, 'static', 'img', 'LogoEEP.png')
    y_logo = margin
    # Se alinean por el centro vertical del bloque de logos, no por el borde superior,
    # para que ambos luzcan balanceados aunque tengan alturas distintas.
    y_logo_ministerio = y_logo + (logo_h_max - logo_ministerio_h) / 2
    y_logo_eep = y_logo + (logo_h_max - logo_eep_h) / 2
    logo_x_izq = margin + logo_inset
    logo_x_der = page_w - margin - logo_inset - logo_eep_w

    if os.path.exists(logo_ministerio):
        pdf.image(logo_ministerio, x=logo_x_izq, y=y_logo_ministerio, w=logo_ministerio_w, h=logo_ministerio_h, type=_tipo_imagen(logo_ministerio))
    if os.path.exists(logo_eep):
        pdf.image(logo_eep, x=logo_x_der, y=y_logo_eep, w=logo_eep_w, h=logo_eep_h, type=_tipo_imagen(logo_eep))

    # El membrete se centra respecto al ancho total de la página (ignorando el
    # ancho/posición de los logos), para que no se desplace si los logos son asimétricos.
    pdf.set_xy(margin, y_logo)
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(usable_w, 5, 'República Bolivariana de Venezuela', ln=2, align='C')
    pdf.set_x(margin)
    pdf.set_font('Arial', 'B', 10)
    pdf.cell(usable_w, 5, 'Ministerio del Poder Popular para la Educación', ln=2, align='C')
    pdf.set_x(margin)
    pdf.cell(usable_w, 5, "E.E.P Daniel O'Leary", ln=2, align='C')
    pdf.set_x(margin)
    pdf.set_font('Arial', '', 9)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(usable_w, 5, 'San Fernando Estado Apure', ln=2, align='C')

    # ── Título ──
    pdf.set_xy(margin, y_logo + logo_h_max + 2)
    pdf.set_font('Arial', 'B', 13)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(usable_w, 8, 'REGISTRO DESCRIPTIVO DE LA ACTUACIÓN DEL ESTUDIANTE', ln=1, align='C', fill=False)

    # ── Datos del Alumno ──
    pdf.ln(2)
    col_w = usable_w / 3
    y_datos = pdf.get_y()

    pdf.set_font('Arial', 'B', 9)
    pdf.set_xy(margin, y_datos)
    pdf.cell(col_w, 6, f"Estudiante: {est.nombre_completo}")
    pdf.set_xy(margin + col_w, y_datos)
    pdf.cell(col_w, 6, f"C.I: {ci_mostrar}")
    pdf.set_xy(margin + 2 * col_w, y_datos)
    pdf.cell(col_w, 6, "Año Escolar: 2026 - 2027")

    # El nombre del grado puede venir combinado con la sección (ej. "Segundo Grado
    # Sección A"); se separan ambos campos para no dejar "Sección:" vacío.
    seccion_fallback = est.literal_escolar or est.literal or '-'
    grado_txt, seccion_txt = _separar_grado_seccion(grado.nombre if grado else None, seccion_fallback)

    y_datos2 = y_datos + 6
    pdf.set_xy(margin, y_datos2)
    pdf.cell(col_w, 6, f"Grado: {grado_txt}")
    pdf.set_xy(margin + col_w, y_datos2)
    pdf.cell(col_w, 6, f"Sección: {seccion_txt}")
    pdf.set_xy(margin + 2 * col_w, y_datos2)
    pdf.cell(col_w, 6, f"Docente(s): {docentes_nombres}")

    pdf.set_draw_color(150, 150, 150)
    y_regla = y_datos2 + 8
    pdf.line(margin, y_regla, page_w - margin, y_regla)

    # ── 3 Columnas: Momento I / II / III ──
    y_top = y_regla + 3
    y_bottom_limit = page_h - 32  # espacio reservado para firmas
    # Anclaje al fondo: "Sugerencias" siempre arranca en esta Y fija (unos mm antes del
    # borde inferior de la columna), sin importar dónde termine el texto de Actuación.
    y_sugerencias = y_bottom_limit - 24
    col_gap = 6
    col_w3 = (usable_w - 2 * col_gap) / 3
    etiquetas = {1: 'I MOMENTO', 2: 'II MOMENTO', 3: 'III MOMENTO'}

    for idx, m in enumerate([1, 2, 3]):
        x = margin + idx * (col_w3 + col_gap)
        x_inner_izq = x + 2
        x_inner_der = x + col_w3 - 2
        proyecto = proyectos.get(m)
        evaluacion = evaluaciones.get(m)

        # Caja de la columna (alto fijo -> las 3 columnas quedan visualmente iguales)
        pdf.set_draw_color(120, 120, 120)
        pdf.rect(x, y_top, col_w3, y_bottom_limit - y_top)

        # Encabezado del Momento (sin relleno: texto negro sobre fondo transparente)
        pdf.set_xy(x, y_top)
        pdf.set_font('Arial', 'B', 10)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(col_w3, 6, etiquetas[m], align='C', fill=False)

        # Título del Proyecto de Aprendizaje + fechas
        pdf.set_xy(x_inner_izq, y_top + 8)
        pdf.set_font('Arial', 'B', 8)
        titulo_pa = proyecto.titulo_proyecto if proyecto and proyecto.titulo_proyecto else 'Proyecto de Aprendizaje'
        pdf.multi_cell(col_w3 - 4, 4, titulo_pa, align='C')

        pdf.set_x(x_inner_izq)
        pdf.set_font('Arial', '', 7)
        fecha_ini = proyecto.fecha_inicio.strftime('%d/%m/%Y') if proyecto and proyecto.fecha_inicio else '--/--/----'
        fecha_fin = proyecto.fecha_cierre.strftime('%d/%m/%Y') if proyecto and proyecto.fecha_cierre else '--/--/----'
        pdf.cell(col_w3 - 4, 5, f"Del {fecha_ini} al {fecha_fin}", align='C', ln=1)

        # Línea divisoria 1: debajo de la Fecha Inicio/Cierre, antes de "ACTUACIÓN DEL ESTUDIANTE"
        y_linea1 = pdf.get_y() + 2
        pdf.set_draw_color(120, 120, 120)
        pdf.line(x_inner_izq, y_linea1, x_inner_der, y_linea1)

        y_cursor = y_linea1 + 3

        # Actuación del Estudiante (Texto Descriptivo)
        pdf.set_xy(x_inner_izq, y_cursor)
        pdf.set_font('Arial', 'B', 7.5)
        pdf.cell(col_w3 - 4, 4, 'ACTUACIÓN DEL ESTUDIANTE:')
        y_cursor = pdf.get_y() + 4

        pdf.set_xy(x_inner_izq, y_cursor)
        pdf.set_font('Arial', '', 8)
        texto_descriptivo = evaluacion.texto_descriptivo if evaluacion and evaluacion.texto_descriptivo else ''
        pdf.multi_cell(col_w3 - 4, 4, texto_descriptivo, align='J')

        # Línea divisoria 2: anclada a y_sugerencias (fija), NO a donde terminó el texto de
        # Actuación -> el espacio sobrante queda atrapado en medio y Sugerencias siempre
        # queda pegado al fondo de la columna.
        y_linea2 = y_sugerencias - 2
        pdf.line(x_inner_izq, y_linea2, x_inner_der, y_linea2)

        # Sugerencias / Recomendaciones (impresas a partir de la coordenada Y fija)
        pdf.set_xy(x_inner_izq, y_sugerencias)
        pdf.set_font('Arial', 'B', 7.5)
        pdf.cell(col_w3 - 4, 4, 'Sugerencias/Recomendaciones:', ln=1)

        pdf.set_x(x_inner_izq)
        pdf.set_font('Arial', 'I', 7.5)
        sugerencias = evaluacion.sugerencias if evaluacion and evaluacion.sugerencias else ''
        pdf.multi_cell(col_w3 - 4, 4, sugerencias, align='L')

    # ── Firmas ──
    y_firma_linea = page_h - 22
    firma_w = usable_w / 4
    etiquetas_firma = ['Directora', 'Enlace Pedagógico', 'Docente', 'Representante']

    pdf.set_draw_color(0, 0, 0)
    pdf.set_font('Arial', '', 8)
    for i, etiqueta in enumerate(etiquetas_firma):
        x = margin + i * firma_w
        pdf.line(x + 8, y_firma_linea, x + firma_w - 8, y_firma_linea)
        pdf.set_xy(x, y_firma_linea + 1)
        pdf.cell(firma_w, 5, etiqueta, align='C')


@academico_bp.route('/generar_boletin_pdf/<int:estudiante_id>')
def generar_boletin_pdf(estudiante_id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    est = Estudiante.query.get_or_404(estudiante_id)
    grado = est.grado

    proyectos = {p.momento: p for p in ProyectoAula.query.filter_by(grado_id=grado.id).all()} if grado else {}
    evaluaciones = {ev.momento: ev for ev in EvaluacionEstudiante.query.filter_by(estudiante_id=est.id).all()}

    pdf = FPDF(orientation='L', unit='mm', format='Letter')
    pdf.set_auto_page_break(auto=False)
    pdf.add_page()
    _dibujar_hoja_boletin(pdf, est, proyectos, evaluaciones)

    nombre_archivo = f"Boletin_de_{est.nombre_completo.strip().replace(' ', '_')}.pdf"
    response = make_response(pdf.output(dest='S').encode('latin1'))
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'inline; filename="{nombre_archivo}"'
    return response


@academico_bp.route('/generar_boletines_masivo')
def generar_boletines_masivo():
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    rol = session.get('nombre_rol')
    if rol not in ['Administrador Supremo', 'Equipo Directivo (Dirección)', 'Docente de Aula']:
        return redirect(url_for('index'))

    grado_id = request.args.get('grado_id')
    grado = None

    if grado_id:
        grado = Grado.query.get_or_404(grado_id)
        if rol == 'Docente de Aula' and session.get('usuario_id') not in [d.id for d in grado.docentes]:
            flash('No tiene acceso a este salón.', 'danger')
            return redirect(url_for('academico.mi_aula'))
    elif rol == 'Docente de Aula':
        grado = Grado.query.filter(Grado.docentes.any(id=session.get('usuario_id'))).first()

    if not grado:
        flash('Debe seleccionar un salón para imprimir los boletines.', 'warning')
        return redirect(url_for('academico.mi_aula'))

    estudiantes = Estudiante.query.filter_by(grado_id=grado.id).order_by(Estudiante.nombre_completo.asc()).all()
    if not estudiantes:
        flash('Este salón no tiene estudiantes registrados.', 'warning')
        return redirect(url_for('academico.mi_aula', grado_id=grado.id))

    # El Proyecto de Aula es por salón (no por estudiante): se consulta una sola vez.
    proyectos = {p.momento: p for p in ProyectoAula.query.filter_by(grado_id=grado.id).all()}

    pdf = FPDF(orientation='L', unit='mm', format='Letter')
    pdf.set_auto_page_break(auto=False)

    for est in estudiantes:
        evaluaciones = {ev.momento: ev for ev in EvaluacionEstudiante.query.filter_by(estudiante_id=est.id).all()}
        pdf.add_page()
        _dibujar_hoja_boletin(pdf, est, proyectos, evaluaciones)

    # El nombre del grado puede venir combinado con la sección (ej. "Segundo Grado
    # Sección A"); se separan para poder nombrar el archivo como Boletines_<Grado>_Seccion_<X>.pdf.
    seccion_fallback = (estudiantes[0].literal_escolar or estudiantes[0].literal or '') if estudiantes else ''
    grado_txt, seccion_txt = _separar_grado_seccion(grado.nombre, seccion_fallback)
    grado_slug = grado_txt.strip().replace(' ', '_')
    if seccion_txt and seccion_txt != '-':
        nombre_archivo = f"Boletines_{grado_slug}_Seccion_{seccion_txt.strip().replace(' ', '_')}.pdf"
    else:
        nombre_archivo = f"Boletines_{grado_slug}.pdf"

    response = make_response(pdf.output(dest='S').encode('latin1'))
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'inline; filename="{nombre_archivo}"'
    return response

# ==========================================
# --- 10. COMPARTIR EXPEDIENTE SEGURO ---
# ==========================================

@academico_bp.route('/solicitar_enlace/<int:estudiante_id>', methods=['POST'])
def solicitar_enlace(estudiante_id):
    if not session.get('logeado'):
        return redirect(url_for('auth.login'))
        
    motivo = request.form.get('motivo')
    if not motivo:
        flash("Debe proporcionar un motivo para la solicitud.", "error")
        return redirect(url_for('academico.mi_aula'))
        
    nueva_solicitud = SolicitudEnlace(
        docente_id=session['usuario_id'],
        estudiante_id=estudiante_id,
        motivo=motivo,
        estado='Pendiente'
    )
    
    db.session.add(nueva_solicitud)
    db.session.commit()
    
    flash("Solicitud de enlace de expediente enviada a Dirección exitosamente.", "success")
    return redirect(url_for('academico.mi_aula'))

@academico_bp.route('/generar_enlace/<int:id>')
def generar_enlace(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    
    import uuid
    estudiante = Estudiante.query.get_or_404(id)
    token = uuid.uuid4().hex
    
    nuevo_enlace = EnlaceTemporal(token=token, estudiante_id=estudiante.id, usado=False)
    db.session.add(nuevo_enlace)
    db.session.commit()
    
    url_acceso = url_for('academico.ver_expediente', token=token, _external=True)
    mensaje = f"Aquí tienes el expediente temporal de {estudiante.nombre_completo}: {url_acceso}\n\n*Nota: Este enlace es de un solo uso y expirará después de abrirlo por razones de seguridad.*"
    
    whatsapp_url = f"https://wa.me/?text={urllib.parse.quote(mensaje)}"
    return redirect(whatsapp_url)

@academico_bp.route('/aprobar_solicitud/<int:id>', methods=['POST', 'GET'])
def aprobar_solicitud(id):
    if session.get('nombre_rol') not in ['Administrador Supremo', 'Equipo Directivo (Dirección)']:
        flash("No tiene permisos para realizar esta acción.", "danger")
        return redirect(url_for('index'))
        
    solicitud = SolicitudEnlace.query.get_or_404(id)
    solicitud.estado = 'Aprobado'
    solicitud.token_generado = uuid.uuid4().hex[:8]
    db.session.commit()
    
    flash(f"Solicitud aprobada exitosamente. Se ha generado un enlace para {solicitud.estudiante.nombre_completo}.", "success")
    return redirect(url_for('index'))

@academico_bp.route('/rechazar_solicitud/<int:id>', methods=['POST', 'GET'])
def rechazar_solicitud(id):
    if session.get('nombre_rol') not in ['Administrador Supremo', 'Equipo Directivo (Dirección)']:
        flash("No tiene permisos para realizar esta acción.", "danger")
        return redirect(url_for('index'))
        
    solicitud = SolicitudEnlace.query.get_or_404(id)
    solicitud.estado = 'Rechazado'
    db.session.commit()
    
    flash("Solicitud rechazada.", "info")
    return redirect(url_for('index'))

@academico_bp.route('/expediente_publico/<token>')
def expediente_publico(token):
    solicitud = SolicitudEnlace.query.filter_by(token_generado=token, estado='Aprobado').first()
    
    if not solicitud:
        return render_template('enlace_expirado.html'), 404
        
    estudiante_data = solicitud.estudiante
    solicitud.estado = 'Utilizado'
    db.session.commit()
        
    # Prevenir que el navegador guarde la página en caché para mayor seguridad
    from flask import make_response
    response = make_response(render_template('expediente_publico.html', estudiante=estudiante_data))
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '-1'
    return response

@academico_bp.route('/ver_expediente/<token>')
def ver_expediente(token):
    enlace = EnlaceTemporal.query.filter_by(token=token).first()
    
    if not enlace or enlace.usado:
        return render_template('enlace_expirado.html'), 403
        
    enlace.usado = True
    db.session.commit()
    
    # Prevenir que el navegador guarde la página en caché
    from flask import make_response
    response = make_response(render_template('ver_expediente_temporal.html', estudiante=enlace.estudiante))
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '-1'
    return response

# ==========================================
# --- 11. GESTIÓN DESCENTRALIZADA (FASE 19) ---
# ==========================================

@academico_bp.route('/nuevo_ingreso_rapido', methods=['POST'])
def nuevo_ingreso_rapido():
    if not session.get('logeado'):
        return redirect(url_for('auth.login'))

    # ── 1. Datos del Representante ──
    cedula_rep = request.form.get('cedula_representante')
    if not cedula_rep:
        flash("Debe indicar la cédula del representante.", "danger")
        return redirect(request.referrer or url_for('academico.mi_aula'))

    representante = Representante.query.filter_by(cedula=cedula_rep).first()
    if not representante:
        representante = Representante(
            cedula=cedula_rep,
            nombre_completo=request.form.get('nombre_representante', ''),
            telefono=request.form.get('telefono_representante'),
            email=request.form.get('email_representante'),
            email_representante=request.form.get('email_representante'),
            parentesco=request.form.get('parentesco'),
            direccion_habitacion=request.form.get('direccion_habitacion'),
            direccion_completa=request.form.get('direccion_completa'),
            ocupacion=request.form.get('ocupacion_representante'),
            lugar_direccion_trabajo=request.form.get('lugar_direccion_trabajo'),
            banco_nombre=request.form.get('banco_nombre'),
            banco_cuenta_numero=request.form.get('banco_cuenta_numero'),
            banco_cuenta_tipo=request.form.get('banco_cuenta_tipo'),
            banco_titular_nombre=request.form.get('banco_titular_nombre'),
            banco_titular_ci=request.form.get('banco_titular_ci')
        )
        db.session.add(representante)
        db.session.flush()
    else:
        # Actualizar datos del representante si vienen nuevos valores
        if request.form.get('nombre_representante'):
            representante.nombre_completo = request.form.get('nombre_representante')
        if request.form.get('telefono_representante'):
            representante.telefono = request.form.get('telefono_representante')
        if request.form.get('email_representante'):
            representante.email = request.form.get('email_representante')
            representante.email_representante = request.form.get('email_representante')
        if request.form.get('parentesco'):
            representante.parentesco = request.form.get('parentesco')
        if request.form.get('direccion_habitacion'):
            representante.direccion_habitacion = request.form.get('direccion_habitacion')
        if request.form.get('direccion_completa'):
            representante.direccion_completa = request.form.get('direccion_completa')
        if request.form.get('ocupacion_representante'):
            representante.ocupacion = request.form.get('ocupacion_representante')
        if request.form.get('lugar_direccion_trabajo'):
            representante.lugar_direccion_trabajo = request.form.get('lugar_direccion_trabajo')
        if request.form.get('banco_nombre'):
            representante.banco_nombre = request.form.get('banco_nombre')
        if request.form.get('banco_cuenta_numero'):
            representante.banco_cuenta_numero = request.form.get('banco_cuenta_numero')
        if request.form.get('banco_cuenta_tipo'):
            representante.banco_cuenta_tipo = request.form.get('banco_cuenta_tipo')
        if request.form.get('banco_titular_nombre'):
            representante.banco_titular_nombre = request.form.get('banco_titular_nombre')
        if request.form.get('banco_titular_ci'):
            representante.banco_titular_ci = request.form.get('banco_titular_ci')
        db.session.flush()

    # ── 2. Generar Cédula Escolar ──
    fecha_nacimiento_str = request.form.get('fecha_nacimiento')
    try:
        fecha_nac = datetime.strptime(fecha_nacimiento_str, '%Y-%m-%d').date()
    except (ValueError, TypeError):
        flash("Fecha de nacimiento inválida.", "danger")
        return redirect(request.referrer or url_for('academico.mi_aula'))

    nro_parto = request.form.get('nro_parto', '1')
    anio_nino = fecha_nac.year
    cedula_escolar_gen = generar_cedula_escolar(nro_parto, anio_nino, cedula_rep)

    # ── 3. Verificar duplicados ──
    estudiante_existente = Estudiante.query.filter_by(cedula_escolar=cedula_escolar_gen).first()
    if estudiante_existente:
        grado_nombre = estudiante_existente.grado.nombre if estudiante_existente.grado else "el sistema"
        flash(f'La solicitud no fue enviada: Este estudiante ya se encuentra registrado en {grado_nombre}.', 'warning')
        return redirect(request.referrer or url_for('academico.mi_aula'))

    # ── 4. Captura exhaustiva de datos del estudiante ──
    talla_val = request.form.get('talla')
    peso_val = request.form.get('peso')
    edad_val = request.form.get('edad')
    nombres_val = request.form.get('nombres', '')
    apellidos_val = request.form.get('apellidos', '')

    nuevo_estudiante = Estudiante(
        # Identidad y Legal
        cedula_escolar=cedula_escolar_gen,
        nombre_completo=f"{nombres_val} {apellidos_val}".strip(),
        nombres=nombres_val,
        apellidos=apellidos_val,
        fecha_nacimiento=fecha_nac,
        lugar_nacimiento=request.form.get('lugar_nacimiento'),
        municipio=request.form.get('municipio'),
        parroquia=request.form.get('parroquia'),
        edad=int(edad_val) if edad_val else None,
        genero=request.form.get('genero'),
        nacionalidad=request.form.get('nacionalidad', 'Venezolana'),
        num_acta=request.form.get('num_acta'),
        num_oficio=request.form.get('num_oficio'),
        direccion_alumno=request.form.get('direccion_alumno'),
        telefono_habitacion=request.form.get('telefono_habitacion'),
        posee_canaima=True if request.form.get('posee_canaima') == 'on' else False,
        # Salud y Antropometría
        talla=float(talla_val) if talla_val else None,
        peso=float(peso_val) if peso_val else None,
        calzado=request.form.get('calzado'),
        talla_camisa=request.form.get('talla_camisa'),
        talla_pantalon=request.form.get('talla_pantalon'),
        tipaje=request.form.get('tipaje'),
        vacunacion_completa=request.form.get('vacunacion_completa'),
        alergias=request.form.get('alergias'),
        neurodivergencia=True if request.form.get('neurodivergencia') == 'on' else False,
        neuro_detalle=request.form.get('neuro_detalle'),
        posee_enfermedad=True if request.form.get('posee_enfermedad') == 'on' else False,
        enfermedad_detalle=request.form.get('enfermedad_detalle'),
        toma_medicamento=True if request.form.get('toma_medicamento') == 'on' else False,
        medicamento_detalle=request.form.get('medicamento_detalle'),
        alergico_medicamento=True if request.form.get('alergico_medicamento') == 'on' else False,
        alergia_medicamento_detalle=request.form.get('alergia_medicamento_detalle'),
        lateralidad=request.form.get('lateralidad'),
        doc_partida=True if request.form.get('doc_partida') == 'on' else False,
        doc_sano=True if request.form.get('doc_sano') == 'on' else False,
        doc_vacuna=True if request.form.get('doc_vacuna') == 'on' else False,
        # Procedencia y Estado Académico
        literal=request.form.get('literal_escolar'),
        literal_escolar=request.form.get('literal_escolar'),
        procedencia=request.form.get('plantel_procedencia'),
        plantel_procedencia=request.form.get('plantel_procedencia'),
        email_estudiante=request.form.get('email_estudiante'),
        es_repetidor=True if request.form.get('es_repetidor') == 'on' else False,
        nuevo_ingreso=True if request.form.get('nuevo_ingreso') == 'on' else True,
        institucion_procedencia=request.form.get('institucion_procedencia'),
        estatus='Activo',
        # Datos de la Madre
        madre_nombre=request.form.get('madre_nombre'),
        madre_ci=request.form.get('madre_ci'),
        madre_ocupacion=request.form.get('madre_ocupacion'),
        madre_telefono=request.form.get('madre_telefono'),
        madre_direccion=request.form.get('madre_direccion'),
        # Datos del Padre
        padre_nombre=request.form.get('padre_nombre'),
        padre_ci=request.form.get('padre_ci'),
        padre_ocupacion=request.form.get('padre_ocupacion'),
        padre_telefono=request.form.get('padre_telefono'),
        padre_direccion=request.form.get('padre_direccion'),
        # Otros y Autorizaciones
        telefono_familiar_extra=request.form.get('telefono_familiar_extra'),
        autorizacion_odontologica=True if request.form.get('autorizacion_odontologica') == 'on' else False,
        # Relaciones
        grado_id=request.form.get('grado_id'),
        representante_id=representante.id
    )

    try:
        db.session.add(nuevo_estudiante)
        db.session.commit()
        flash(f"Estudiante {nuevo_estudiante.nombre_completo} registrado exitosamente.", "success")
    except Exception as e:
        db.session.rollback()
        flash('Ocurrió un error inesperado al guardar el estudiante. Inténtalo de nuevo.', 'danger')

    return redirect(request.referrer or url_for('academico.mi_aula'))

@academico_bp.route('/actualizar_estudiante_rapido/<int:id>', methods=['POST'])
def actualizar_estudiante_rapido(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    
    estudiante = Estudiante.query.get_or_404(id)
    
    # Crear la solicitud en lugar de un UPDATE directo
    nueva_solicitud = SolicitudActualizacion(
        docente_id=session['usuario_id'],
        estudiante_id=estudiante.id,
        rep_telefono=request.form.get('rep_telefono'),
        rep_direccion=request.form.get('rep_direccion'),
        alergias=request.form.get('alergias'),
        neuro_detalle=request.form.get('neuro_detalle'),
        estado='Pendiente'
    )
    
    db.session.add(nueva_solicitud)
    db.session.commit()
    
    flash(f"Solicitud de actualización para {estudiante.nombre_completo} enviada a Dirección para su revisión.", "info")
    return redirect(request.referrer or url_for('academico.mi_aula'))

@academico_bp.route('/actualizar_medidas/<int:id>', methods=['POST'])
def actualizar_medidas(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))

    estudiante = Estudiante.query.get_or_404(id)
    grado_id = request.form.get('grado_id')

    def _a_float(valor):
        try:
            return float(valor) if valor not in (None, '') else None
        except ValueError:
            return None

    # Guardado directo: a diferencia de "Actualizar" (contacto/ficha médica),
    # las medidas físicas no pasan por aprobación de Dirección.
    estudiante.peso = _a_float(request.form.get('peso'))
    estudiante.talla = _a_float(request.form.get('talla'))
    estudiante.talla_camisa = request.form.get('talla_camisa')
    estudiante.talla_pantalon = request.form.get('talla_pantalon')
    estudiante.calzado = request.form.get('calzado')

    db.session.commit()
    flash(f"Medidas de {estudiante.nombre_completo} actualizadas correctamente.", "success")

    url = url_for('academico.mi_aula')
    if session.get('nombre_rol') in ['Administrador Supremo', 'Equipo Directivo (Dirección)'] and grado_id:
        url = url_for('academico.mi_aula', grado_id=grado_id)
    return redirect(url)

@academico_bp.route('/aprobar_actualizacion/<int:id>', methods=['POST', 'GET'])
def aprobar_actualizacion(id):
    if session.get('nombre_rol') not in ['Administrador Supremo', 'Equipo Directivo (Dirección)']:
        flash("No tiene permisos para realizar esta acción.", "danger")
        return redirect(url_for('index'))
        
    solicitud = SolicitudActualizacion.query.get_or_404(id)
    estudiante = solicitud.estudiante
    
    if estudiante.representante_info:
        if solicitud.rep_telefono:
            estudiante.representante_info.telefono = solicitud.rep_telefono
        if solicitud.rep_direccion:
            estudiante.representante_info.direccion_completa = solicitud.rep_direccion
            
    if solicitud.alergias is not None:
        estudiante.alergias = solicitud.alergias
    if solicitud.neuro_detalle is not None:
        estudiante.neuro_detalle = solicitud.neuro_detalle
        estudiante.neurodivergencia = bool(solicitud.neuro_detalle.strip())
        
    solicitud.estado = 'Aprobado'
    db.session.commit()
    
    flash(f"Actualización para {estudiante.nombre_completo} aprobada y aplicada.", "success")
    return redirect(url_for('index'))

@academico_bp.route('/rechazar_actualizacion/<int:id>', methods=['POST', 'GET'])
def rechazar_actualizacion(id):
    if session.get('nombre_rol') not in ['Administrador Supremo', 'Equipo Directivo (Dirección)']:
        flash("No tiene permisos para realizar esta acción.", "danger")
        return redirect(url_for('index'))
        
    solicitud = SolicitudActualizacion.query.get_or_404(id)
    solicitud.estado = 'Rechazado'
    db.session.commit()
    
    flash("Solicitud de actualización rechazada.", "info")
    return redirect(url_for('index'))

# ==========================================
# --- ASISTENCIA POR ESTUDIANTE Y ALERTAS ---
# ==========================================

@academico_bp.route('/asistencia_estudiantes', methods=['GET'])
def asistencia_estudiantes():
    """Página principal de asistencia por estudiante."""
    if not session.get('logeado'):
        return redirect(url_for('auth.login'))
    
    grados = Grado.query.all()
    fecha_hoy = date.today().strftime('%Y-%m-%d')
    
    return render_template('asistencia_estudiantes.html', grados=grados, fecha_hoy=fecha_hoy)

@academico_bp.route('/api/asistencia_estudiantes/<int:grado_id>/<string:fecha>', methods=['GET'])
def api_asistencia_estudiantes(grado_id, fecha):
    """API: Devuelve la lista de estudiantes de un grado con su estatus de asistencia para una fecha."""
    if not session.get('logeado'):
        return jsonify({'error': 'No autorizado'}), 401
    
    try:
        fecha_obj = datetime.strptime(fecha, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Formato de fecha inválido'}), 400
    
    estudiantes = Estudiante.query.filter_by(grado_id=grado_id, estatus='Activo').order_by(Estudiante.nombre_completo.asc()).all()
    
    resultado = []
    for est in estudiantes:
        # Buscar si ya tiene asistencia registrada para esta fecha
        asistencia = AsistenciaEstudiante.query.filter_by(
            estudiante_id=est.id, fecha=fecha_obj
        ).first()
        
        resultado.append({
            'id': est.id,
            'nombre_completo': est.nombre_completo,
            'cedula_escolar': est.cedula_escolar,
            'genero': est.genero or 'N/A',
            'estatus': asistencia.estatus if asistencia else None  # None = no registrado aún
        })
    
    return jsonify({'estudiantes': resultado, 'total': len(resultado)})

@academico_bp.route('/guardar_asistencia_estudiantes', methods=['POST'])
def guardar_asistencia_estudiantes():
    """Guarda la asistencia de todos los estudiantes de un grado y ejecuta el disparador de alertas."""
    if not session.get('logeado'):
        return jsonify({'error': 'No autorizado'}), 401
    
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No se recibieron datos'}), 400
    
    grado_id = data.get('grado_id')
    fecha_str = data.get('fecha')
    asistencias = data.get('asistencias', [])  # [{estudiante_id: X, estatus: 'Presente'|'Ausente'|'Justificado'}]
    
    if not grado_id or not fecha_str or not asistencias:
        return jsonify({'error': 'Faltan datos requeridos'}), 400
    
    try:
        fecha_obj = datetime.strptime(fecha_str, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Formato de fecha inválido'}), 400
    
    try:
        # Insertar o actualizar asistencia de cada estudiante
        for registro in asistencias:
            est_id = registro.get('estudiante_id')
            estatus = registro.get('estatus', 'Presente')
            
            # Validar estatus
            if estatus not in ('Presente', 'Ausente', 'Justificado'):
                estatus = 'Presente'
            
            # Buscar registro existente para esta fecha y estudiante
            existente = AsistenciaEstudiante.query.filter_by(
                estudiante_id=est_id, fecha=fecha_obj
            ).first()
            
            if existente:
                existente.estatus = estatus
                existente.grado_id = grado_id
            else:
                nuevo = AsistenciaEstudiante(
                    estudiante_id=est_id,
                    grado_id=grado_id,
                    fecha=fecha_obj,
                    estatus=estatus
                )
                db.session.add(nuevo)
        
        db.session.flush()  # Aplicar cambios antes de verificar alertas
        
        # =============================================
        # DISPARADOR: Verificar alertas de Defensoría
        # =============================================
        alertas_generadas = 0
        
        # Calcular rango de la semana (lunes a viernes) de la fecha dada
        dia_semana = fecha_obj.weekday()  # 0=Lunes, 6=Domingo
        lunes = fecha_obj - timedelta(days=dia_semana)
        viernes = lunes + timedelta(days=4)
        semana_iso = fecha_obj.strftime('%G-W%V')  # Ej: '2026-W31'
        
        # Solo verificar estudiantes marcados como Ausentes
        ausentes_ids = [r['estudiante_id'] for r in asistencias if r.get('estatus') == 'Ausente']
        
        for est_id in ausentes_ids:
            # Contar inasistencias en la semana actual
            conteo_ausencias = AsistenciaEstudiante.query.filter(
                AsistenciaEstudiante.estudiante_id == est_id,
                AsistenciaEstudiante.estatus == 'Ausente',
                AsistenciaEstudiante.fecha >= lunes,
                AsistenciaEstudiante.fecha <= viernes
            ).count()
            
            if conteo_ausencias >= 3:
                # Verificar si ya existe alerta pendiente para esta semana
                alerta_existente = AlertaDefensoria.query.filter_by(
                    estudiante_id=est_id,
                    semana_iso=semana_iso
                ).first()
                
                if not alerta_existente:
                    estudiante = Estudiante.query.get(est_id)
                    nombre = estudiante.nombre_completo if estudiante else f'ID {est_id}'
                    grado = Grado.query.get(grado_id)
                    grado_nombre = grado.nombre if grado else ''
                    
                    # Escalación automática: si el estudiante ya tiene 2+ alertas previas
                    # de semanas distintas, escalar directamente a Visita Domiciliaria
                    alertas_previas = AlertaDefensoria.query.filter(
                        AlertaDefensoria.estudiante_id == est_id,
                        AlertaDefensoria.semana_iso != semana_iso
                    ).count()
                    
                    if alertas_previas >= 2:
                        estatus_auto = 'Visita Domiciliaria'
                        motivo_texto = f'⚠️ ESCALADA AUTOMÁTICA: El estudiante {nombre} ({grado_nombre}) acumula {conteo_ausencias} inasistencias en la semana del {lunes.strftime("%d/%m/%Y")} al {viernes.strftime("%d/%m/%Y")}. Este estudiante tiene {alertas_previas} alertas previas de semanas anteriores.'
                    else:
                        estatus_auto = 'Pendiente'
                        motivo_texto = f'El estudiante {nombre} ({grado_nombre}) acumula {conteo_ausencias} inasistencias en la semana del {lunes.strftime("%d/%m/%Y")} al {viernes.strftime("%d/%m/%Y")}.'
                    
                    nueva_alerta = AlertaDefensoria(
                        estudiante_id=est_id,
                        fecha_emision=date.today(),
                        motivo=motivo_texto,
                        estatus_atencion=estatus_auto,
                        semana_iso=semana_iso
                    )
                    db.session.add(nueva_alerta)
                    alertas_generadas += 1
        
        # =============================================
        # ACTUALIZAR REGISTRO DIARIO (AsistenciaDiaria)
        # =============================================
        grado = Grado.query.get(grado_id)
        if grado:
            varones_asist = 0
            hembras_asist = 0
            
            for registro in asistencias:
                if registro.get('estatus') in ('Presente', 'Justificado'):
                    est = Estudiante.query.get(registro.get('estudiante_id'))
                    if est:
                        genero_est = (est.genero or '').upper()
                        if genero_est.startswith('M') or genero_est == 'VARON' or genero_est == 'NIÑO':
                            varones_asist += 1
                        else:
                            hembras_asist += 1
                            
            mat_total = grado.total_varones + grado.total_hembras
            asist_total = varones_asist + hembras_asist
            porc = round((asist_total / mat_total) * 100, 2) if mat_total > 0 else 0
            
            reg_diario = AsistenciaDiaria.query.filter_by(fecha=fecha_obj, grado_seccion=grado.nombre).first()
            if reg_diario:
                reg_diario.varones = varones_asist
                reg_diario.hembras = hembras_asist
                reg_diario.asistentes = asist_total
                reg_diario.porcentaje = porc
            else:
                nuevo_diario = AsistenciaDiaria(
                    fecha=fecha_obj,
                    grado_seccion=grado.nombre,
                    matricula_total=mat_total,
                    varones=varones_asist,
                    hembras=hembras_asist,
                    asistentes=asist_total,
                    porcentaje=porc,
                    usuario_id=session.get('usuario_id', 1)
                )
                db.session.add(nuevo_diario)

        db.session.commit()
        
        mensaje = f'Asistencia guardada correctamente para {len(asistencias)} estudiantes.'
        if alertas_generadas > 0:
            mensaje += f' ⚠️ Se generaron {alertas_generadas} alerta(s) para Defensoría.'
        
        return jsonify({'success': True, 'message': mensaje, 'alertas_generadas': alertas_generadas})
    
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Error al guardar: {str(e)}'}), 500

@academico_bp.route('/alertas_defensoria')
def alertas_defensoria():
    """Vista de alertas generadas para Defensoría."""
    if not session.get('logeado'):
        return redirect(url_for('auth.login'))
    
    # Solo Defensoría, Directivo o Admin pueden ver
    rol = session.get('nombre_rol', '')
    if rol not in ['Defensoría Estudiantil', 'Equipo Directivo (Dirección)', 'Administrador Supremo']:
        flash('No tienes permisos para acceder a esta sección.', 'error')
        return redirect(url_for('index'))
    
    alertas = AlertaDefensoria.query.order_by(AlertaDefensoria.fecha_emision.desc()).all()
    return render_template('alertas_defensoria.html', alertas=alertas)

@academico_bp.route('/actualizar_alerta/<int:id>', methods=['POST'])
def actualizar_alerta(id):
    """Actualiza el estatus de atención de una alerta."""
    if not session.get('logeado'):
        return jsonify({'error': 'No autorizado'}), 401
    
    alerta = AlertaDefensoria.query.get_or_404(id)
    nuevo_estatus = request.form.get('estatus_atencion', request.json.get('estatus_atencion') if request.is_json else None)
    
    if nuevo_estatus in ('Pendiente', 'Contactado', 'Visita Domiciliaria'):
        alerta.estatus_atencion = nuevo_estatus
        db.session.commit()
        # No usamos flash() aquí para evitar que los mensajes se acumulen y salgan en otros módulos
    
    referer = request.headers.get("Referer")
    if referer and "defensoria" in referer and "alertas_defensoria" not in referer:
        return redirect(url_for('defensoria'))
    return redirect(url_for('academico.alertas_defensoria'))

@academico_bp.route('/enterado_alerta/<int:id>', methods=['POST'])
def enterado_alerta(id):
    if not session.get('logeado'): return redirect(url_for('auth.login'))
    rol = session.get('nombre_rol')
    if rol not in ['Administrador Supremo', 'Equipo Directivo (Dirección)', 'Docente de Aula']:
        return "Acceso Denegado", 403
    alerta = AlertaDefensoria.query.get_or_404(id)
    # Simplemente eliminamos la alerta de incidencia para que ya no aparezca
    db.session.delete(alerta)
    db.session.commit()
    
    grado_id = request.args.get('grado_id')
    return redirect(url_for('academico.mi_aula', grado_id=grado_id))

@academico_bp.route('/eliminar_alerta/<int:id>', methods=['POST'])
def eliminar_alerta(id):
    """Elimina una alerta de defensoría permanentemente."""
    if not session.get('logeado'):
        return jsonify({'error': 'No autorizado'}), 401
    
    # Solo Defensoría, Directivo o Admin pueden borrar
    rol = session.get('nombre_rol', '')
    if rol not in ['Defensoría Estudiantil', 'Equipo Directivo (Dirección)', 'Administrador Supremo']:
        return redirect(url_for('index'))
        
    alerta = AlertaDefensoria.query.get_or_404(id)
    db.session.delete(alerta)
    db.session.commit()
    # No usamos flash() para evitar mensajes huérfanos en portal_trabajador
    
    # Redirect back to where they came from (defensoria or alertas_defensoria)
    referer = request.headers.get("Referer")
    if referer and "defensoria" in referer and "alertas_defensoria" not in referer:
        return redirect(url_for('defensoria'))
    return redirect(url_for('academico.alertas_defensoria'))
